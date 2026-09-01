from __future__ import annotations

import hashlib
import io
import json
import os
import re
import secrets
import shutil
import subprocess
import zipfile
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from openpyxl import load_workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from ..audit import audit
from ..models import (
    PlanCheckAssumption,
    PlanCheckCase,
    PlanCheckDocument,
    PlanCheckGate,
    PlanCheckRevision,
    WorkspaceDocument,
)

MAX_FILE_SIZE = 20 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".docx", ".xlsx", ".ifc", ".dwg"}
REQUIRED_CATEGORIES = ("site_plan", "floor_plan", "elevations", "sections", "technical_description")
GATE_KEYS = ("input", "engineering", "commercial", "finance", "executive")
GATE_ROLES = {
    "input": {"technical-prep", "project-manager", "platform-admin"},
    "engineering": {"designer", "technical-prep", "platform-admin"},
    "commercial": {"sales", "managing-director", "platform-admin"},
    "finance": {"finance", "managing-director", "platform-admin"},
    "executive": {"owner", "managing-director", "platform-admin"},
}
FINAL_ROLES = {"owner", "managing-director", "platform-admin"}
RUNTIME_ROOT = Path(os.getenv("PLATFORM_RUNTIME_ROOT", "/app/runtime")) / "plancheck"


def utcnow() -> datetime:
    return datetime.now(UTC)


def _json(value: str, fallback: Any) -> Any:
    try:
        return json.loads(value)
    except (TypeError, json.JSONDecodeError):
        return fallback


def _identity(user: object) -> tuple[str, str]:
    return str(getattr(user, "role", "")), str(getattr(user, "email", "")).strip().lower()


def _case(db: Session, case_id: str) -> PlanCheckCase:
    row = db.scalar(select(PlanCheckCase).where(PlanCheckCase.case_id == case_id))
    if row is None:
        raise KeyError(case_id)
    return row


def _revision(db: Session, revision_id: str) -> PlanCheckRevision:
    row = db.scalar(select(PlanCheckRevision).where(PlanCheckRevision.revision_id == revision_id))
    if row is None:
        raise KeyError(revision_id)
    return row


def _current_revision(db: Session, case: PlanCheckCase) -> PlanCheckRevision:
    return _revision(db, case.current_revision_id)


def _rows_for_revision(
    db: Session, revision_id: str
) -> tuple[list[PlanCheckDocument], list[PlanCheckAssumption], list[PlanCheckGate]]:
    documents = list(
        db.scalars(
            select(PlanCheckDocument)
            .where(PlanCheckDocument.revision_id == revision_id)
            .order_by(PlanCheckDocument.created_at)
        )
    )
    assumptions = list(
        db.scalars(
            select(PlanCheckAssumption)
            .where(PlanCheckAssumption.revision_id == revision_id)
            .order_by(PlanCheckAssumption.created_at)
        )
    )
    gates = list(
        db.scalars(
            select(PlanCheckGate)
            .where(PlanCheckGate.revision_id == revision_id)
            .order_by(PlanCheckGate.id)
        )
    )
    return documents, assumptions, gates


def _confidence(documents: list[PlanCheckDocument]) -> tuple[str, list[str]]:
    present = {row.category for row in documents if row.validation_status == "verified"}
    missing = [category for category in REQUIRED_CATEGORIES if category not in present]
    count = len(REQUIRED_CATEGORIES) - len(missing)
    return ("A" if count == 5 else "B" if count == 4 else "C" if count >= 2 else "D"), missing


def _refresh_snapshot(db: Session, revision: PlanCheckRevision) -> None:
    documents, assumptions, gates = _rows_for_revision(db, revision.revision_id)
    confidence, missing = _confidence(documents)
    payload = {
        "input": _json(revision.input_json, {}),
        "documents": [
            {"category": row.category, "sha256": row.content_sha256, "name": row.file_name}
            for row in documents
        ],
        "assumptions": [
            {
                "id": row.assumption_id,
                "impact": row.impact,
                "status": row.status,
                "description": row.description,
            }
            for row in assumptions
        ],
        "gates": [
            {"key": row.gate_key, "decision": row.decision, "actor": row.decided_by}
            for row in gates
        ],
    }
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    revision.snapshot_sha256 = hashlib.sha256(canonical.encode("utf-8")).hexdigest()
    revision.confidence_class = confidence
    revision.missing_items_json = json.dumps(missing, ensure_ascii=False)
    revision.final_eligible = (
        confidence in {"A", "B"}
        and not any(row.impact == "high" and row.status == "open" for row in assumptions)
        and len(gates) == len(GATE_KEYS)
        and all(row.decision == "approved" for row in gates)
    )


def _emit_event(
    db: Session,
    *,
    case: PlanCheckCase,
    revision: PlanCheckRevision,
    event_type: str,
    status: str,
    actor: str,
    evidence_url: str | None = None,
) -> None:
    from ..schemas import EventIn
    from .integration import ingest_event

    next_action = (
        "A tervcsomag feltöltése és a hiányok pótlása."
        if event_type == "PLANCHECK_CASE_CREATED"
        else "A PlanCheck eredmény átvezetése a projektfolyamatokba."
    )
    ingest_event(
        db,
        EventIn(
            event_id=f"EVT-PLC-{uuid4().hex[:14].upper()}",
            dedupe_key=f"{event_type}:{case.case_id}:v{revision.version}",
            project_id=case.project_id,
            source_module="plancheck",
            event_type=event_type,
            object_type="PlanCheckCase",
            object_id=case.case_id,
            status=status,
            responsible="Műszaki előkészítés",
            next_action=next_action,
            evidence_url=evidence_url,
            payload={
                "summary": f"PlanCheck {case.case_id} v{revision.version}: {status}",
                "confidence_class": revision.confidence_class,
                "snapshot_sha256": revision.snapshot_sha256,
            },
            route_to=["crm", "my-imperial", "buildconfig"],
        ),
        actor=actor,
    )


def _new_revision(db: Session, case: PlanCheckCase, actor: str) -> PlanCheckRevision:
    if case.status in {"sendable", "not_sendable"}:
        raise ValueError("Lezárt PlanCheck ügy csak külön, auditált újranyitással módosítható.")
    previous = _current_revision(db, case)
    old_documents, old_assumptions, _old_gates = _rows_for_revision(db, previous.revision_id)
    version = case.current_revision + 1
    revision = PlanCheckRevision(
        revision_id=f"PCR-{uuid4().hex[:16].upper()}",
        case_id=case.case_id,
        version=version,
        input_json=previous.input_json,
        snapshot_sha256="0" * 64,
        confidence_class=previous.confidence_class,
        missing_items_json=previous.missing_items_json,
        created_by=actor,
    )
    db.add(revision)
    db.flush()
    for document_row in old_documents:
        db.add(
            PlanCheckDocument(
                document_id=f"PCD-{uuid4().hex[:16].upper()}",
                revision_id=revision.revision_id,
                category=document_row.category,
                file_name=document_row.file_name,
                mime_type=document_row.mime_type,
                extension=document_row.extension,
                file_size=document_row.file_size,
                content_sha256=document_row.content_sha256,
                storage_path=document_row.storage_path,
                page_count=document_row.page_count,
                extracted_text=document_row.extracted_text,
                validation_status=document_row.validation_status,
                uploaded_by=document_row.uploaded_by,
            )
        )
    for assumption_row in old_assumptions:
        db.add(
            PlanCheckAssumption(
                assumption_id=f"EA-{uuid4().hex[:14].upper()}",
                revision_id=revision.revision_id,
                description=assumption_row.description,
                impact=assumption_row.impact,
                owner=assumption_row.owner,
                status=assumption_row.status,
                resolution=assumption_row.resolution,
                resolved_by=assumption_row.resolved_by,
                resolved_at=assumption_row.resolved_at,
                created_by=assumption_row.created_by,
            )
        )
    for gate_key in GATE_KEYS:
        db.add(PlanCheckGate(revision_id=revision.revision_id, gate_key=gate_key))
    case.current_revision = version
    case.current_revision_id = revision.revision_id
    case.status = "intake"
    case.final_report_document_id = None
    case.finalized_by = None
    case.finalized_at = None
    db.flush()
    return revision


def create_case(
    db: Session,
    *,
    project_id: str,
    title: str,
    contact_name: str,
    contact_email: str,
    intake: dict[str, Any],
    actor: str,
) -> tuple[dict[str, Any], str]:
    if (
        len(project_id.strip()) < 3
        or len(title.strip()) < 3
        or len(contact_name.strip()) < 2
        or "@" not in contact_email
    ):
        raise ValueError("A ProjectID, az ügy neve és az érvényes kapcsolattartó adatai kötelezők.")
    token = secrets.token_urlsafe(32)
    case_id = f"PLC-{uuid4().hex[:14].upper()}"
    revision_id = f"PCR-{uuid4().hex[:16].upper()}"
    case = PlanCheckCase(
        case_id=case_id,
        project_id=project_id.strip(),
        title=title.strip(),
        contact_name=contact_name.strip(),
        contact_email=contact_email.strip().lower(),
        current_revision_id=revision_id,
        upload_token_hash=hashlib.sha256(token.encode()).hexdigest(),
        upload_token_expires_at=utcnow() + timedelta(days=30),
        created_by=actor.lower(),
    )
    revision = PlanCheckRevision(
        revision_id=revision_id,
        case_id=case_id,
        version=1,
        input_json=json.dumps(intake, ensure_ascii=False, sort_keys=True),
        snapshot_sha256="0" * 64,
        confidence_class="D",
        missing_items_json=json.dumps(REQUIRED_CATEGORIES),
        created_by=actor.lower(),
    )
    db.add_all([case, revision])
    db.flush()
    for gate_key in GATE_KEYS:
        db.add(PlanCheckGate(revision_id=revision_id, gate_key=gate_key))
    db.flush()
    _refresh_snapshot(db, revision)
    audit(
        db,
        actor=actor,
        action="plancheck.case.created",
        entity_type="plancheck_case",
        entity_id=case_id,
        after={"project_id": case.project_id, "revision_id": revision_id},
    )
    _emit_event(
        db,
        case=case,
        revision=revision,
        event_type="PLANCHECK_CASE_CREATED",
        status="intake",
        actor=actor,
    )
    return case_detail(db, case_id), token


def list_cases(db: Session) -> list[dict[str, Any]]:
    return [
        case_detail(db, row.case_id)
        for row in db.scalars(select(PlanCheckCase).order_by(PlanCheckCase.updated_at.desc()))
    ]


def case_detail(db: Session, case_id: str) -> dict[str, Any]:
    case = _case(db, case_id)
    revision = _current_revision(db, case)
    documents, assumptions, gates = _rows_for_revision(db, revision.revision_id)
    upload_expires_at = case.upload_token_expires_at
    if upload_expires_at.tzinfo is None:
        upload_expires_at = upload_expires_at.replace(tzinfo=UTC)
    return {
        "case": case,
        "revision": revision,
        "input": _json(revision.input_json, {}),
        "missing_items": _json(revision.missing_items_json, []),
        "documents": documents,
        "assumptions": assumptions,
        "gates": gates,
        "high_open_assumptions": sum(
            row.impact == "high" and row.status == "open" for row in assumptions
        ),
        "upload_link_active": (
            case.status not in {"sendable", "not_sendable"}
            and upload_expires_at > utcnow()
        ),
    }


def rotate_upload_link(
    db: Session, case_id: str, actor: str, *, valid_days: int = 30
) -> tuple[dict[str, Any], str]:
    """Replace the customer upload credential without ever exposing the prior token."""
    if valid_days < 1 or valid_days > 90:
        raise ValueError("A feltöltési hivatkozás érvényessége 1 és 90 nap közötti lehet.")
    case = _case(db, case_id)
    if case.status in {"sendable", "not_sendable"}:
        raise ValueError("Lezárt PlanCheck ügyhöz nem adható új feltöltési hivatkozás.")
    token = secrets.token_urlsafe(32)
    previous_expiry = case.upload_token_expires_at
    case.upload_token_hash = hashlib.sha256(token.encode()).hexdigest()
    case.upload_token_expires_at = utcnow() + timedelta(days=valid_days)
    audit(
        db,
        actor=actor,
        action="plancheck.upload_link.rotated",
        entity_type="plancheck_case",
        entity_id=case.case_id,
        before={"expires_at": previous_expiry.isoformat()},
        after={"expires_at": case.upload_token_expires_at.isoformat()},
    )
    db.commit()
    return case_detail(db, case_id), token


def revoke_upload_link(db: Session, case_id: str, actor: str) -> dict[str, Any]:
    """Fail closed immediately and invalidate any previously distributed upload URL."""
    case = _case(db, case_id)
    previous_expiry = case.upload_token_expires_at
    case.upload_token_hash = hashlib.sha256(secrets.token_bytes(48)).hexdigest()
    case.upload_token_expires_at = utcnow() - timedelta(seconds=1)
    audit(
        db,
        actor=actor,
        action="plancheck.upload_link.revoked",
        entity_type="plancheck_case",
        entity_id=case.case_id,
        before={"expires_at": previous_expiry.isoformat()},
        after={"expires_at": case.upload_token_expires_at.isoformat()},
    )
    db.commit()
    return case_detail(db, case_id)


def case_for_token(db: Session, token: str) -> PlanCheckCase:
    digest = hashlib.sha256(token.encode()).hexdigest()
    case = db.scalar(select(PlanCheckCase).where(PlanCheckCase.upload_token_hash == digest))
    expires_at = None
    if case is not None:
        expires_at = case.upload_token_expires_at
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=UTC)
    if case is None or expires_at is None or expires_at < utcnow():
        raise PermissionError("A feltöltési hivatkozás érvénytelen vagy lejárt.")
    if case.status in {"sendable", "not_sendable"}:
        raise ValueError("A lezárt PlanCheck ügy feltöltése le van tiltva.")
    return case


def _validate_file(file_name: str, content: bytes) -> tuple[str, int | None, str | None]:
    extension = Path(file_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Nem támogatott fájlformátum.")
    if not content or len(content) > MAX_FILE_SIZE:
        raise ValueError("A fájl mérete 1 bájt és 20 MB között lehet.")
    if content.startswith((b"MZ", b"\x7fELF")) or b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE" in content:
        raise ValueError("A fájl végrehajtható vagy kártevőgyanús tartalmat tartalmaz.")
    signatures = {
        ".pdf": content.startswith(b"%PDF-"),
        ".jpg": content.startswith(b"\xff\xd8\xff"),
        ".jpeg": content.startswith(b"\xff\xd8\xff"),
        ".png": content.startswith(b"\x89PNG\r\n\x1a\n"),
        ".docx": content.startswith(b"PK"),
        ".xlsx": content.startswith(b"PK"),
        ".ifc": content.lstrip().startswith(b"ISO-10303-21"),
        ".dwg": content.startswith(b"AC10"),
    }
    if not signatures[extension]:
        raise ValueError("A fájl tartalma nem egyezik a kiterjesztéssel.")
    page_count: int | None = None
    extracted: str | None = None
    if extension in {".docx", ".xlsx"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                if (
                    "[Content_Types].xml" not in names
                    or (
                        extension == ".docx" and not any(name.startswith("word/") for name in names)
                    )
                    or (extension == ".xlsx" and not any(name.startswith("xl/") for name in names))
                ):
                    raise ValueError
        except (zipfile.BadZipFile, ValueError) as exc:
            raise ValueError("A feltöltött Office-fájl szerkezete érvénytelen.") from exc
    if extension == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            extracted = "\n".join((page.extract_text() or "") for page in reader.pages)[:200_000]
        except Exception as exc:
            raise ValueError("A PDF nem olvasható vagy sérült.") from exc
    elif extension == ".docx":
        extracted = "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs)[:200_000]
    elif extension == ".xlsx":
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        values: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                values.extend(str(value) for value in row if value is not None)
                if sum(map(len, values)) > 200_000:
                    break
        extracted = "\n".join(values)[:200_000]
    elif extension == ".ifc":
        extracted = content.decode("utf-8", errors="replace")[:200_000]
    scanner = shutil.which("clamscan")
    if scanner:
        result = subprocess.run(
            [scanner, "--no-summary", "-"],
            input=content,
            capture_output=True,
            timeout=60,
            check=False,
        )
        if result.returncode != 0:
            raise ValueError("A kártevő-ellenőrzés elutasította a fájlt.")
    elif os.getenv("PLANCHECK_REQUIRE_MALWARE_SCANNER", "false").lower() == "true":
        raise ValueError("A kötelező kártevő-ellenőrző szolgáltatás nem érhető el.")
    return extension, page_count, extracted


def upload_document(
    db: Session,
    *,
    token: str,
    category: str,
    file_name: str,
    mime_type: str,
    content: bytes,
    uploader: str,
) -> dict[str, Any]:
    if category not in REQUIRED_CATEGORIES and category != "other":
        raise ValueError("Ismeretlen tervdokumentum-kategória.")
    case = case_for_token(db, token)
    extension, page_count, extracted = _validate_file(file_name, content)
    digest = hashlib.sha256(content).hexdigest()
    revision = _new_revision(db, case, uploader)
    directory = RUNTIME_ROOT / case.case_id / revision.revision_id
    directory.mkdir(parents=True, exist_ok=False)
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(file_name).name)[:180]
    path = directory / f"{digest[:16]}-{safe_name}"
    path.write_bytes(content)
    document = PlanCheckDocument(
        document_id=f"PCD-{uuid4().hex[:16].upper()}",
        revision_id=revision.revision_id,
        category=category,
        file_name=Path(file_name).name,
        mime_type=mime_type or "application/octet-stream",
        extension=extension,
        file_size=len(content),
        content_sha256=digest,
        storage_path=str(path),
        page_count=page_count,
        extracted_text=extracted,
        uploaded_by=uploader.lower(),
    )
    db.add(document)
    db.flush()
    _refresh_snapshot(db, revision)
    audit(
        db,
        actor=uploader,
        action="plancheck.document.uploaded",
        entity_type="plancheck_revision",
        entity_id=revision.revision_id,
        after={"document_id": document.document_id, "sha256": digest, "version": revision.version},
    )
    db.commit()
    return case_detail(db, case.case_id)


def add_assumption(
    db: Session, case_id: str, *, description: str, impact: str, owner: str, actor: str
) -> dict[str, Any]:
    if impact not in {"low", "medium", "high"} or len(description.strip()) < 10 or "@" not in owner:
        raise ValueError("Az EA leírása, hatása és felelős e-mail-címe kötelező.")
    case = _case(db, case_id)
    revision = _new_revision(db, case, actor)
    db.add(
        PlanCheckAssumption(
            assumption_id=f"EA-{uuid4().hex[:14].upper()}",
            revision_id=revision.revision_id,
            description=description.strip(),
            impact=impact,
            owner=owner.lower(),
            created_by=actor.lower(),
        )
    )
    db.flush()
    _refresh_snapshot(db, revision)
    audit(
        db,
        actor=actor,
        action="plancheck.assumption.created",
        entity_type="plancheck_revision",
        entity_id=revision.revision_id,
        after={"impact": impact, "version": revision.version},
    )
    db.commit()
    return case_detail(db, case_id)


def resolve_assumption(
    db: Session, case_id: str, assumption_id: str, *, resolution: str, actor: str
) -> dict[str, Any]:
    if len(resolution.strip()) < 10:
        raise ValueError("A feloldás indoklása legalább 10 karakter.")
    case = _case(db, case_id)
    current = _current_revision(db, case)
    source = db.scalar(
        select(PlanCheckAssumption).where(
            PlanCheckAssumption.revision_id == current.revision_id,
            PlanCheckAssumption.assumption_id == assumption_id,
        )
    )
    if source is None or source.status != "open":
        raise KeyError(assumption_id)
    revision = _new_revision(db, case, actor)
    target = db.scalar(
        select(PlanCheckAssumption).where(
            PlanCheckAssumption.revision_id == revision.revision_id,
            PlanCheckAssumption.description == source.description,
            PlanCheckAssumption.owner == source.owner,
        )
    )
    if target is None:
        raise RuntimeError("A feltételezés revíziós másolata hiányzik.")
    target.status = "resolved"
    target.resolution = resolution.strip()
    target.resolved_by = actor.lower()
    target.resolved_at = utcnow()
    _refresh_snapshot(db, revision)
    audit(
        db,
        actor=actor,
        action="plancheck.assumption.resolved",
        entity_type="plancheck_revision",
        entity_id=revision.revision_id,
        after={"assumption_id": target.assumption_id, "version": revision.version},
    )
    db.commit()
    return case_detail(db, case_id)


def submit_review(db: Session, case_id: str, actor: str) -> dict[str, Any]:
    case = _case(db, case_id)
    revision = _current_revision(db, case)
    documents, _assumptions, _gates = _rows_for_revision(db, revision.revision_id)
    if not documents:
        raise ValueError("Ellenőrzéshez legalább egy hiteles tervdokumentum szükséges.")
    case.status = "review"
    audit(
        db,
        actor=actor,
        action="plancheck.review.submitted",
        entity_type="plancheck_case",
        entity_id=case_id,
        after={"revision_id": revision.revision_id},
    )
    db.commit()
    return case_detail(db, case_id)


def review_gate(
    db: Session, case_id: str, *, gate_key: str, decision: str, note: str, user: object
) -> dict[str, Any]:
    role, email = _identity(user)
    if gate_key not in GATE_ROLES or role not in GATE_ROLES[gate_key]:
        raise PermissionError("Ehhez a PlanCheck kapuhoz nincs jogosultsága.")
    if decision not in {"approve", "reject"} or len(note.strip()) < 10:
        raise ValueError("A döntés és legalább 10 karakteres indoklás kötelező.")
    case = _case(db, case_id)
    if case.status != "review":
        raise ValueError("Kapudöntés csak ellenőrzés alatt álló ügyön rögzíthető.")
    revision = _current_revision(db, case)
    gate = db.scalar(
        select(PlanCheckGate).where(
            PlanCheckGate.revision_id == revision.revision_id, PlanCheckGate.gate_key == gate_key
        )
    )
    if gate is None or gate.decision != "pending":
        raise ValueError("A PlanCheck kapu már lezárult vagy nem található.")
    used = {
        row.decided_by
        for row in db.scalars(
            select(PlanCheckGate).where(PlanCheckGate.revision_id == revision.revision_id)
        )
        if row.decided_by
    }
    if email == case.created_by or email in used:
        raise ValueError("A készítő és az egyes kapuk jóváhagyói külön személyek kell legyenek.")
    gate.decision = "approved" if decision == "approve" else "rejected"
    gate.note = note.strip()
    gate.decided_by = email
    gate.decided_at = utcnow()
    _refresh_snapshot(db, revision)
    audit(
        db,
        actor=email,
        action="plancheck.gate.decided",
        entity_type="plancheck_gate",
        entity_id=f"{revision.revision_id}:{gate_key}",
        after={"decision": gate.decision, "note": gate.note},
    )
    db.commit()
    return case_detail(db, case_id)


def _report(
    case: PlanCheckCase,
    revision: PlanCheckRevision,
    documents: list[PlanCheckDocument],
    assumptions: list[PlanCheckAssumption],
    gates: list[PlanCheckGate],
    outcome: str,
) -> tuple[Path, str]:
    directory = RUNTIME_ROOT / case.case_id / "reports"
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"PlanCheck-{case.case_id}-v{revision.version}-{outcome}.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    y = 800
    for line in [
        "IMPERIAL INTELLIGENCE - PlanCheck v0.1",
        f"Ugy: {case.case_id}",
        f"ProjectID: {case.project_id}",
        f"Revízió: {revision.version}",
        f"Eredmény: {outcome.upper()}",
        f"Bizalmi osztály: {revision.confidence_class}",
        f"Dokumentumok: {len(documents)}",
        f"Nyitott magas hatású feltételezések: {sum(a.impact == 'high' and a.status == 'open' for a in assumptions)}",
        "Kapuk: " + ", ".join(f"{g.gate_key}={g.decision}" for g in gates),
        f"Snapshot SHA-256: {revision.snapshot_sha256}",
    ]:
        pdf.drawString(40, y, line)
        y -= 24
    pdf.save()
    return path, hashlib.sha256(path.read_bytes()).hexdigest()


def finalize_case(
    db: Session, case_id: str, *, outcome: str, note: str, user: object
) -> dict[str, Any]:
    role, email = _identity(user)
    if role not in FINAL_ROLES:
        raise PermissionError("A PlanCheck végső döntéshez nincs jogosultsága.")
    if outcome not in {"sendable", "not_sendable"} or len(note.strip()) < 10:
        raise ValueError("A végső döntés és indoklás kötelező.")
    case = _case(db, case_id)
    if case.status != "review":
        raise ValueError("Csak ellenőrzés alatt álló ügy zárható le.")
    revision = _current_revision(db, case)
    documents, assumptions, gates = _rows_for_revision(db, revision.revision_id)
    _refresh_snapshot(db, revision)
    if outcome == "sendable" and not revision.final_eligible:
        raise ValueError(
            "SENDABLE csak A/B bizalmi osztállyal, öt jóváhagyott kapuval és nyitott magas hatású feltételezés nélkül adható."
        )
    path, digest = _report(case, revision, documents, assumptions, gates, outcome)
    document_id = f"DOC-PLC-{uuid4().hex[:12].upper()}"
    db.add(
        WorkspaceDocument(
            document_id=document_id,
            project_id=case.project_id,
            title=f"PlanCheck {outcome.upper()} – {case.case_id}",
            category="plancheck_report",
            source_system="plancheck",
            source_url=f"file://{path}",
            mime_type="application/pdf",
            version_label=f"v{revision.version}",
            approval_status="approved",
            verification_status="sha256_verified",
            confidentiality="internal",
            owner="Műszaki előkészítés",
            extracted_summary=f"{outcome.upper()}; confidence={revision.confidence_class}; SHA-256={digest}",
            metadata_json=json.dumps(
                {
                    "sha256": digest,
                    "snapshot_sha256": revision.snapshot_sha256,
                    "local_path": str(path),
                },
                ensure_ascii=False,
            ),
        )
    )
    case.status = outcome
    case.final_report_document_id = document_id
    case.finalized_by = email
    case.finalized_at = utcnow()
    audit(
        db,
        actor=email,
        action="plancheck.finalized",
        entity_type="plancheck_case",
        entity_id=case_id,
        after={"outcome": outcome, "report_document_id": document_id, "note": note.strip()},
    )
    _emit_event(
        db,
        case=case,
        revision=revision,
        event_type="PLANCHECK_FINALIZED",
        status=outcome,
        actor=email,
        evidence_url=f"document://{document_id}",
    )
    return case_detail(db, case_id)
