from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timedelta, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document as create_docx_document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from ..audit import audit
from ..models import ProjectRegistry, TaskRecord, WorkspaceDocument
from ..schemas import WorkspaceDocumentIn
from .workspace import create_document


BASE_DIR = Path(__file__).resolve().parents[1]
CATALOG_PATH = BASE_DIR / "canonical_documents" / "templates.json"
OUTPUT_ROOT = BASE_DIR.parent / "runtime" / "canonical_documents"
CANONICAL_FOLDER_DRIVE_ID = "1WDIvn93b-3aBNtGllN2npdtpmL21ibNi"
CANONICAL_REGISTRY_DRIVE_FILE_ID = "1-UE60l3OBUSt9qxlRgyIW9ShOUc2WqZtVYAQIU9TVrQ"

CATEGORY_LABELS = {
    "CARE": "Imperial Care és garancia",
    "CLOSE": "Projektzárás",
    "COM": "Hivatalos kommunikáció",
    "COMM": "Üzembe helyezés",
    "CXM": "Ügyfélélmény",
    "DES": "Tervezés és tervellenőrzés",
    "FIN": "Pénzügy és könyvelés",
    "HRA": "HR és munkajog",
    "HSE": "Munkavédelem és biztonság",
    "LEG": "Jogi iratok",
    "OPS": "Projekt- és helyszíni operáció",
    "PART": "Partner- és alvállalkozó-kezelés",
    "PER": "Engedélyezés és hatóság",
    "PROC": "Beszerzés és tender",
    "QA": "Minőségbiztosítás",
    "SAL": "Értékesítés és projektindítás",
    "SCH": "Ütemezés és határidő",
    "SEL": "Műszaki választások",
}

MANDATORY_CORE = {
    "TPL-OPS-019", "TPL-DES-001", "TPL-DES-003", "TPL-PER-002",
    "TPL-OPS-020", "TPL-OPS-002", "TPL-OPS-001", "TPL-OPS-005",
    "TPL-CLOSE-001", "TPL-CLOSE-002", "TPL-FIN-004", "TPL-HRA-003",
    "TPL-HRA-004",
}

EVENT_CONTROL = {
    "TPL-OPS-003": "Változás vagy pótmunka",
    "TPL-QA-001": "Műszaki kérdés",
    "TPL-DES-002": "Tervütközés vagy szakági koordináció",
    "TPL-DES-004": "Gyártmányterv vagy egyedi gyártás",
    "TPL-QA-002": "Takarás előtti vagy visszatartási pont",
    "TPL-QA-003": "Nemmegfelelőség",
    "TPL-HSE-002": "Veszélyes munkavégzés",
    "TPL-HSE-003": "Baleset vagy majdnem-baleset",
    "TPL-OPS-017": "Munkafelfüggesztés vagy újraindítás",
    "TPL-SCH-002": "Határidőt érintő esemény",
    "TPL-SCH-003": "Igazolt késés és helyreállítási igény",
    "TPL-CARE-001": "Garanciális igény bejelentése",
    "TPL-CARE-002": "Garanciális kivizsgálás",
    "TPL-CARE-003": "Garanciális javítás lezárása",
    "TPL-PART-002": "Partneri incidens vagy szerződésszegés",
    "TPL-PART-003": "Partner újramegbízási vagy kizárási döntés",
    "TPL-FIN-001": "Számla-, TIG- vagy szerződéseltérés",
    "TPL-FIN-002": "Előleg vagy rendkívüli kifizetés",
    "TPL-FIN-003": "Visszatartás vagy felszabadítás",
    "TPL-COM-001": "Hiánypótlási igény",
    "TPL-COM-002": "Akadály vagy késedelem hivatalos közlése",
    "TPL-COM-003": "Döntés vagy jóváhagyás kérése",
    "TPL-COM-004": "Ügyféloldali bemenet vagy döntés hiánya",
}

ROLE_CATEGORIES = {
    "project-manager": {"CARE", "CLOSE", "COM", "COMM", "CXM", "DES", "HSE", "OPS", "PART", "PER", "PROC", "QA", "SAL", "SCH", "SEL"},
    "finance": {"FIN", "CLOSE", "PROC"},
    "legal": {"LEG", "COM", "FIN"},
    "technical-prep": {"DES", "OPS", "PER", "QA", "SCH", "SEL"},
    "designer": {"DES", "PER", "QA", "SCH", "SEL"},
    "subcontractor": {"HSE", "OPS", "PART", "QA"},
    "sales": {"COM", "SAL"},
}


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _template_id(title: str) -> str:
    match = re.match(r"^(TPL-[A-Z]+-\d+[A-Z]?)", title)
    if not match:
        raise ValueError(f"Érvénytelen kanonikus TemplateID: {title}")
    return match.group(1)


def _category(template_id: str) -> str:
    return template_id.split("-")[1]


def _version(title: str, content: str) -> str:
    match = re.search(r"(?:Verzió|Version)\s*:\s*([^\r\n]+)", content, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    match = re.search(r"_v(\d+(?:\.\d+)?)", title, re.IGNORECASE)
    return f"v{match.group(1)}" if match else "v1.0"


def _display_name(title: str, template_id: str) -> str:
    value = title[len(template_id):].lstrip("_").replace("_", " ")
    value = re.sub(r"\s+kanonikus\s+v\d+(?:\.\d+)?(?:\s+ÉRVÉNYES)?$", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\s+v\d+(?:\.\d+)?\s+ÉRVÉNYES$", "", value, flags=re.IGNORECASE)
    return value.strip(" -")


def _purpose(content: str) -> str:
    match = re.search(r"(?:^|\n)Cél\s*:\s*([^\r\n]+)", content, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    return next((line for line in lines if len(line) > 35 and not line.upper().startswith("IMPERIAL")), "Kanonikus működési iratminta.")


def _legal_level(content: str) -> str | None:
    match = re.search(r"(?:^|\n)Szint\s*:\s*([^\r\n]+)", content, re.IGNORECASE)
    return match.group(1).strip() if match else None


@lru_cache(maxsize=1)
def _raw_catalog() -> dict[str, Any]:
    return json.loads(CATALOG_PATH.read_text(encoding="utf-8-sig"))


@lru_cache(maxsize=1)
def canonical_template_catalog() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for source in _raw_catalog()["templates"]:
        template_id = _template_id(source["title"])
        category = _category(template_id)
        content = source["content"].lstrip("\ufeff")
        tier = "A" if template_id in MANDATORY_CORE else "B" if template_id in EVENT_CONTROL else "C"
        rows.append({
            **source,
            "content": content,
            "template_id": template_id,
            "category": category,
            "category_label": CATEGORY_LABELS.get(category, category),
            "display_name": _display_name(source["title"], template_id),
            "version": _version(source["title"], content),
            "purpose": _purpose(content),
            "tier": tier,
            "tier_label": {"A": "Kötelező magirat", "B": "Eseményalapú kontrollirat", "C": "Támogató operatív irat"}[tier],
            "trigger": EVENT_CONTROL.get(template_id),
            "legal_level": _legal_level(content),
            "sha256": hashlib.sha256(content.encode("utf-8")).hexdigest(),
            "source_verified": bool(content and source.get("id") and source.get("url")),
        })
    return rows


def role_can_use_template(role: str, template: dict[str, Any]) -> bool:
    if role in {"owner", "managing-director", "platform-admin"}:
        return True
    return template["category"] in ROLE_CATEGORIES.get(role, set())


def list_canonical_templates(*, role: str, category: str | None = None, query: str | None = None) -> list[dict[str, Any]]:
    needle = (query or "").strip().casefold()
    rows = []
    for row in canonical_template_catalog():
        if not role_can_use_template(role, row):
            continue
        if category and row["category"] != category:
            continue
        if needle and needle not in f'{row["template_id"]} {row["display_name"]} {row["purpose"]} {row["category_label"]}'.casefold():
            continue
        rows.append(row)
    return rows


def get_canonical_template(template_id: str, *, role: str) -> dict[str, Any]:
    row = next((item for item in canonical_template_catalog() if item["template_id"] == template_id), None)
    if not row:
        raise KeyError(template_id)
    if not role_can_use_template(role, row):
        raise PermissionError(template_id)
    return row


def canonical_template_status() -> dict[str, Any]:
    rows = canonical_template_catalog()
    return {
        "healthy": len(rows) == 86 and all(row["source_verified"] for row in rows),
        "template_count": len(rows),
        "verified_count": sum(1 for row in rows if row["source_verified"]),
        "category_count": len({row["category"] for row in rows}),
        "folder_drive_id": CANONICAL_FOLDER_DRIVE_ID,
        "registry_drive_file_id": CANONICAL_REGISTRY_DRIVE_FILE_ID,
        "captured_at": _raw_catalog()["source"]["captured_at"],
    }


def _safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_") or "document"


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    """Apply the standard_business_brief fixed-DXA table contract."""
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = None
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:type"), "dxa")
        node.set(qn("w:w"), str(value))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))


def _set_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Oldal ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)


def _apply_business_brief_styles(document: DocxDocument) -> None:
    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, (46, 116, 181)),
        ("Heading 2", 13, 12, 6, (46, 116, 181)),
        ("Heading 3", 12, 8, 4, (31, 77, 120)),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def _render_docx(template: dict[str, Any], instance: dict[str, Any], output_path: Path) -> str:
    document = create_docx_document()
    _apply_business_brief_styles(document)
    title = document.add_heading(template["display_name"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f'{template["template_id"]} · {template["version"]} · {template["tier_label"]}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    _set_page_field(footer)
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(102, 112, 128)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    labels = {
        "document_id": "DocumentID", "project_id": "ProjectID", "related_object_id": "Kapcsolódó objektum",
        "trigger_reason": "Kiváltó esemény", "occurred_at": "Esemény időpontja", "owner": "Felelős",
        "participants": "Résztvevők", "facts": "Tények és körülmények", "decision": "Döntés / eredmény",
        "actions": "Intézkedések és határidők", "evidence_ids": "Bizonyítékok / EvidenceID-k", "notes": "Megjegyzések",
    }
    for key, label in labels.items():
        value = instance.get(key)
        if value:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True
    _set_table_geometry(table, [2700, 6660])
    document.add_heading("Kanonikus forrásszöveg", level=1)
    for line in template["content"].splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith(("•", "–")):
            document.add_paragraph(clean[1:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", clean) and len(clean) < 150:
            document.add_paragraph(clean, style="Heading 2")
        elif clean.isupper() and len(clean) < 150:
            document.add_paragraph(clean, style="Heading 2")
        else:
            document.add_paragraph(clean)
    document.add_heading("Kitöltési és jóváhagyási nyom", level=1)
    document.add_paragraph(
        "A fenti kanonikus forrásszöveg változatlan pillanatképe. A dokumentumpéldányhoz tartozó kitöltött adatok, "
        "jóváhagyások, bizonyítékok és státuszok az Imperial Intelligence Document & Evidence rekordjában kezelendők."
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return hashlib.sha256(output_path.read_bytes()).hexdigest()


def _duplicate_instance(db: Session, *, template_id: str, project_id: str | None, related_object_id: str | None) -> WorkspaceDocument | None:
    rows = db.scalars(select(WorkspaceDocument).where(
        WorkspaceDocument.source_system == "canonical_document_generator",
        WorkspaceDocument.project_id == project_id,
        WorkspaceDocument.approval_status.in_(["draft", "pending_review", "approved"]),
    )).all()
    for row in rows:
        metadata = json.loads(row.metadata_json or "{}")
        if metadata.get("template_id") == template_id and metadata.get("related_object_id") == related_object_id:
            return row
    return None


def instantiate_canonical_template(
    db: Session,
    *,
    template_id: str,
    role: str,
    actor: str,
    owner: str,
    project_id: str | None = None,
    related_object_id: str | None = None,
    trigger_reason: str | None = None,
    occurred_at: str | None = None,
    participants: str | None = None,
    facts: str | None = None,
    decision: str | None = None,
    actions: str | None = None,
    evidence_ids: str | None = None,
    notes: str | None = None,
    due_at: str | None = None,
) -> dict[str, Any]:
    template = get_canonical_template(template_id, role=role)
    if project_id and not db.scalar(select(ProjectRegistry).where(ProjectRegistry.project_id == project_id)):
        raise ValueError("A megadott ProjectID nem létezik.")
    if template["tier"] == "B" and not (trigger_reason or "").strip():
        raise ValueError("Eseményalapú kontrolliratnál a kiváltó esemény megadása kötelező.")
    if template["tier"] == "B" and not (related_object_id or "").strip():
        raise ValueError("Eseményalapú kontrolliratnál a kapcsolódó esemény vagy objektum azonosítója kötelező.")
    duplicate = _duplicate_instance(db, template_id=template_id, project_id=project_id, related_object_id=related_object_id)
    if duplicate:
        raise ValueError(f"Ehhez az eseményhez már létezik elsődleges irat: {duplicate.document_id}.")

    values = {
        "document_id": None,
        "project_id": project_id,
        "related_object_id": (related_object_id or "").strip() or None,
        "trigger_reason": (trigger_reason or template.get("trigger") or "").strip() or None,
        "occurred_at": (occurred_at or "").strip() or None,
        "owner": owner,
        "participants": (participants or "").strip() or None,
        "facts": (facts or "").strip() or None,
        "decision": (decision or "").strip() or None,
        "actions": (actions or "").strip() or None,
        "evidence_ids": (evidence_ids or "").strip() or None,
        "notes": (notes or "").strip() or None,
    }
    needs_legal = bool(template.get("legal_level") and re.search(r"L[123]", template["legal_level"]))
    row = create_document(db, WorkspaceDocumentIn(
        title=template["display_name"],
        project_id=project_id,
        category=f'canonical_{template["category"].lower()}',
        source_system="canonical_document_generator",
        source_url=None,
        drive_file_id=template["id"],
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        version_label=template["version"],
        approval_status="pending_review" if needs_legal else "draft",
        verification_status="verified",
        confidentiality="confidential" if template["category"] in {"HRA", "LEG"} else "internal",
        owner=owner,
        extracted_summary=f'{template["template_id"]} · {template["purpose"]}',
        metadata={
            **values,
            "template_id": template["template_id"],
            "template_version": template["version"],
            "template_tier": template["tier"],
            "template_source_drive_id": template["id"],
            "template_source_url": template["url"],
            "template_source_sha256": template["sha256"],
            "canonical_source_immutable": True,
            "legal_review_level": template.get("legal_level"),
        },
    ), actor=actor)
    values["document_id"] = row.document_id
    row.title = f'{template["display_name"]} · {row.document_id}'
    output_path = OUTPUT_ROOT / _safe_name(project_id or "general") / f"{row.document_id}_{_safe_name(template_id)}.docx"
    artifact_sha = _render_docx(template, values, output_path)
    metadata = json.loads(row.metadata_json)
    metadata.update({"document_id": row.document_id, "artifact_sha256": artifact_sha, "local_path": str(output_path)})
    row.metadata_json = json.dumps(metadata, ensure_ascii=False)
    row.source_url = f"/documents/files/{row.document_id}"

    now = _utcnow()
    try:
        final_due = datetime.fromisoformat(due_at).replace(tzinfo=timezone.utc) if due_at else now + timedelta(days=5)
    except ValueError:
        final_due = now + timedelta(days=5)
    task_specs = [
        ("data", "Adatok és bizonyítékok teljességi ellenőrzése", now + timedelta(days=1), "high"),
        ("approval", "Irat szakmai és jogosultsági jóváhagyása", now + timedelta(days=2), "high"),
        ("closure", "Irat lezárása, aláírása és archiválása", final_due, "normal"),
    ]
    task_ids: list[str] = []
    for key, title, task_due, priority in task_specs:
        task_id = f"TASK-DOC-{hashlib.sha256(f'{row.document_id}:{key}'.encode()).hexdigest()[:14].upper()}"
        task_ids.append(task_id)
        db.add(TaskRecord(
            task_id=task_id,
            project_id=project_id or "GENERAL",
            source_event_id=row.document_id,
            title=f"{title} · {template_id}",
            description=f'{template["tier_label"]}. DocumentID: {row.document_id}. A kanonikus forrás nem írható felül.',
            assignee=owner if "@" in owner else None,
            due_at=task_due,
            priority=priority,
            status="open",
            executive_relevance=needs_legal or template["tier"] == "A",
        ))
    audit(db, actor=actor, action="canonical_document.instantiated", entity_type="document", entity_id=row.document_id, after={
        "template_id": template_id, "project_id": project_id, "related_object_id": related_object_id,
        "source_sha256": template["sha256"], "artifact_sha256": artifact_sha, "task_ids": task_ids,
    })
    db.commit()
    db.refresh(row)
    return {"document": row, "template": template, "task_ids": task_ids, "artifact_sha256": artifact_sha}
