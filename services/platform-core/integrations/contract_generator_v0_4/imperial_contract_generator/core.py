from __future__ import annotations

import argparse
import copy
import hashlib
import json
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from datetime import UTC, date, datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

from docx import Document as create_docx_document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

GENERATOR_VERSION = "0.4.0"
OWNER_POLICY_VERSION = "ICG-PAY-2026-07-18"
DESIGNER_PAYMENT_TERM_DAYS = 8
EXECUTION_SUBCONTRACTOR_PAYMENT_TERM_DAYS = 30


class ContractValidationError(ValueError):
    """Raised when a contract package does not satisfy a blocking rule."""


CONTRACT_TYPES = {
    "customer_type_house_design_build",
    "customer_construction",
    "customer_design_execution_plans",
    "subcontractor_design",
    "subcontractor_execution",
}
CUSTOMER_TYPES = {
    "customer_type_house_design_build",
    "customer_construction",
    "customer_design_execution_plans",
}
SUBCONTRACTOR_TYPES = {"subcontractor_design", "subcontractor_execution"}
EXECUTION_SUBCONTRACTOR_TYPE = "subcontractor_execution"
DESIGN_SUBCONTRACTOR_TYPE = "subcontractor_design"
REQUIRED_IDS = ("CompanyID", "PersonID", "OpportunityID", "ProjectID", "PartnerID")

LEGAL_SOURCE_NOTES = {
    "invoice_authenticity": "Áfa tv. 168/A. § (1)–(2): hitelesség, sértetlenség, olvashatóság és megbízható ellenőrzési kapcsolat.",
    "invoice_content": "Áfa tv. 169. §: a számla kötelező adattartalma.",
    "invoice_correction": "Áfa tv. 168. § (2) és 170. §: a számla módosítása számlával egy tekintet alá eső okirattal.",
}


@dataclass(frozen=True)
class ValidationIssue:
    code: str
    message: str
    blocking: bool = True

    def as_dict(self) -> dict[str, Any]:
        return {"code": self.code, "message": self.message, "blocking": self.blocking}


def money(value: Any) -> Decimal:
    return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def pct_amount(base: Any, percent: Any) -> Decimal:
    return (money(base) * Decimal(str(percent)) / Decimal("100")).quantize(
        Decimal("0.01"), rounding=ROUND_HALF_UP
    )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _valid_https_url(value: str) -> bool:
    try:
        parsed = urlparse(value)
        return parsed.scheme == "https" and bool(parsed.netloc)
    except Exception:
        return False


def _valid_email(value: str) -> bool:
    return bool(re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", value or ""))


def _valid_iso_date(value: Any) -> bool:
    try:
        date.fromisoformat(str(value))
        return True
    except (TypeError, ValueError):
        return False


def _get(data: dict[str, Any], dotted: str) -> Any:
    cur: Any = data
    for part in dotted.split("."):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(part)
    return cur


def _missing(data: dict[str, Any], fields: Iterable[str]) -> list[str]:
    result: list[str] = []
    for field in fields:
        value = _get(data, field)
        if value is None or (isinstance(value, str) and not value.strip()) or value == []:
            result.append(field)
    return result


def select_contract_type(data: dict[str, Any]) -> str:
    relationship = data.get("relationship")
    service = data.get("service")
    if relationship == "partner" and service == "design":
        return DESIGN_SUBCONTRACTOR_TYPE
    if relationship == "partner" and service == "construction":
        return EXECUTION_SUBCONTRACTOR_TYPE
    if relationship == "customer" and service == "construction":
        return "customer_construction"
    if relationship == "customer" and service == "design" and data.get("execution_plans"):
        return "customer_design_execution_plans"
    if relationship == "customer" and service == "design_build" and data.get("type_house"):
        return "customer_type_house_design_build"
    raise ContractValidationError("No approved contract type matches the CRM opportunity.")


def validate_party(data: dict[str, Any], prefix: str, company_only: bool = False) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    party = _get(data, prefix) or {}
    party_type = party.get("party_type")
    if company_only and party_type != "company":
        issues.append(ValidationIssue("PARTY_TYPE_INVALID", f"{prefix}.party_type must be company."))
        return issues
    if party_type not in {"company", "natural_person", "sole_proprietor"}:
        issues.append(ValidationIssue("PARTY_TYPE_MISSING", f"Mandatory or invalid party_type: {prefix}.party_type."))
        return issues

    common = ["name", "address", "postal_address", "email", "phone"]
    for field in _missing(data, [f"{prefix}.{x}" for x in common]):
        issues.append(ValidationIssue("PARTY_DATA_MISSING", f"Mandatory party field missing: {field}."))

    if party_type in {"company", "sole_proprietor"}:
        fields = [
            "short_name", "registration_number", "tax_number", "registered_office",
            "bank_account", "representative", "representative_title",
        ]
        for field in _missing(data, [f"{prefix}.{x}" for x in fields]):
            issues.append(ValidationIssue("COMPANY_DATA_MISSING", f"Mandatory company field missing: {field}."))
    else:
        fields = ["birth_place", "birth_date", "mother_name", "identity_document_type", "identity_document_number"]
        for field in _missing(data, [f"{prefix}.{x}" for x in fields]):
            issues.append(ValidationIssue("PERSONAL_DATA_MISSING", f"Mandatory personal identification field missing: {field}."))
        if party.get("birth_date") and not _valid_iso_date(party.get("birth_date")):
            issues.append(ValidationIssue("BIRTH_DATE_INVALID", f"{prefix}.birth_date must use ISO format YYYY-MM-DD."))

    if party.get("email") and not _valid_email(str(party.get("email"))):
        issues.append(ValidationIssue("EMAIL_INVALID", f"Invalid email address: {prefix}.email."))
    return issues


def validate_commercial(data: dict[str, Any]) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    required = [
        "commercial.net_price", "commercial.vat_percent", "commercial.vat_amount",
        "commercial.gross_price", "commercial.currency", "commercial.payment_schedule",
    ]
    for field in _missing(data, required):
        issues.append(ValidationIssue("COMMERCIAL_DATA_MISSING", f"Mandatory commercial field missing: {field}."))
    c = data.get("commercial", {})
    try:
        net = money(c.get("net_price", 0))
        vat_pct = Decimal(str(c.get("vat_percent", 0)))
        vat = money(c.get("vat_amount", 0))
        gross = money(c.get("gross_price", 0))
        expected_vat = pct_amount(net, vat_pct)
        if net <= 0 or gross <= 0:
            issues.append(ValidationIssue("PRICE_INVALID", "Net and gross contract price must be greater than zero."))
        if abs(vat - expected_vat) > Decimal("1.00"):
            issues.append(ValidationIssue("VAT_CALCULATION_MISMATCH", "VAT amount does not match net price and VAT percentage."))
        if abs(gross - (net + vat)) > Decimal("1.00"):
            issues.append(ValidationIssue("GROSS_CALCULATION_MISMATCH", "Gross price does not equal net price plus VAT."))
    except Exception:
        issues.append(ValidationIssue("PRICE_FORMAT_INVALID", "Contract price fields must be numeric."))

    schedule = c.get("payment_schedule", [])
    if schedule:
        total_pct = Decimal("0")
        for idx, row in enumerate(schedule, 1):
            if not row.get("milestone") or row.get("percent") is None or not row.get("due_rule"):
                issues.append(ValidationIssue("PAYMENT_SCHEDULE_ROW_INVALID", f"Payment schedule row {idx} is incomplete."))
            else:
                total_pct += Decimal(str(row.get("percent")))
        if total_pct != Decimal("100"):
            issues.append(ValidationIssue("PAYMENT_SCHEDULE_TOTAL_INVALID", "Payment schedule percentages must total 100%."))
    return issues


def validate_contract(data: dict[str, Any]) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    contract_type = data.get("contract_type")
    if contract_type not in CONTRACT_TYPES:
        issues.append(ValidationIssue("TYPE_INVALID", "Unknown or missing contract_type."))

    for field in _missing(data, ["contract_number", "contract_date", "contract_place"]):
        issues.append(ValidationIssue("CONTRACT_HEADER_MISSING", f"Mandatory contract header field missing: {field}."))
    if data.get("contract_date") and not _valid_iso_date(data.get("contract_date")):
        issues.append(ValidationIssue("CONTRACT_DATE_INVALID", "Contract date must use ISO format YYYY-MM-DD."))

    ids = data.get("ids", {})
    for field in REQUIRED_IDS:
        if not ids.get(field):
            issues.append(ValidationIssue("ID_MISSING", f"Mandatory identifier missing: {field}."))

    issues.extend(validate_party(data, "internal_entity", company_only=True))
    issues.extend(validate_party(data, "counterparty"))

    project_required = ["project.name", "project.site_address", "project.parcel_number", "project.scope"]
    if contract_type in {"customer_type_house_design_build", "customer_design_execution_plans", "subcontractor_design"}:
        project_required.append("project.gross_floor_area_m2")
    if contract_type == DESIGN_SUBCONTRACTOR_TYPE:
        project_required.append("project.procedure_type")
    for field in _missing(data, project_required):
        issues.append(ValidationIssue("PROJECT_DATA_MISSING", f"Mandatory project field missing: {field}."))

    schedule_fields = ["schedule.start_date", "schedule.deadline"]
    if contract_type in {"customer_construction", EXECUTION_SUBCONTRACTOR_TYPE}:
        schedule_fields.append("schedule.site_handover_date")
    for field in _missing(data, schedule_fields):
        issues.append(ValidationIssue("SCHEDULE_DATA_MISSING", f"Mandatory schedule field missing: {field}."))
    for field in schedule_fields:
        value = _get(data, field)
        if value and not _valid_iso_date(value):
            issues.append(ValidationIssue("SCHEDULE_DATE_INVALID", f"Invalid ISO date: {field}."))

    issues.extend(validate_commercial(data))

    delivery = data.get("delivery_requirements", {})
    if delivery.get("postal_required") is not True:
        issues.append(ValidationIssue("POSTAL_DELIVERY_RULE_MISSING", "Trackable postal delivery of the signed original must be mandatory."))
    if delivery.get("electronic_required") is not True:
        issues.append(ValidationIssue("ELECTRONIC_DELIVERY_RULE_MISSING", "Electronic delivery of the signed contract must be mandatory."))
    if delivery.get("postal_method") not in {"registered_return_receipt", "courier_proof_of_delivery", "trackable_mail"}:
        issues.append(ValidationIssue("POSTAL_METHOD_INVALID", "A trackable postal/courier method is mandatory."))
    if int(delivery.get("postal_original_copies", 0) or 0) < 1:
        issues.append(ValidationIssue("POSTAL_ORIGINAL_COPY_MISSING", "At least one signed original contract copy must be sent to the partner."))
    if delivery.get("electronic_format") not in {"SIGNED_PDF", "QUALIFIED_E_SIGNATURE", "ADVANCED_E_SIGNATURE"}:
        issues.append(ValidationIssue("ELECTRONIC_FORMAT_INVALID", "The mandatory electronic signed-copy format must be defined."))

    attachments = data.get("attachments", [])
    required_attachments = set(data.get("required_attachments", []))
    present = {item.get("type") for item in attachments if item.get("status") in {"APPROVED", "SIGNED", "FINAL"}}
    for item in sorted(required_attachments - present):
        issues.append(ValidationIssue("ANNEX_MISSING", f"Mandatory annex missing or not approved: {item}."))

    if contract_type == EXECUTION_SUBCONTRACTOR_TYPE:
        issues.extend(validate_execution_subcontractor_controls(data))
    elif contract_type == DESIGN_SUBCONTRACTOR_TYPE:
        for field in _missing(data, [
            "designer_controls.professional_liability_insurance.insurer",
            "designer_controls.professional_liability_insurance.policy_number",
            "designer_controls.professional_liability_insurance.coverage_amount",
            "schedule.milestones", "invoice_controls.payment_term_days",
        ]):
            issues.append(ValidationIssue("DESIGN_CONTROL_MISSING", f"Mandatory designer field missing: {field}."))
        design_payment_days = int(data.get("invoice_controls", {}).get("payment_term_days", 0) or 0)
        if design_payment_days != DESIGNER_PAYMENT_TERM_DAYS:
            issues.append(ValidationIssue("DESIGN_PAYMENT_TERM_INVALID", f"Designer invoice payment term must be exactly {DESIGNER_PAYMENT_TERM_DAYS} calendar days."))
        if data.get("subcontractor_controls"):
            issues.append(ValidationIssue(
                "CONSTRUCTION_CONTROLS_NOT_APPLICABLE",
                "Construction-only retention, upstream TIG, service-deduction and defect-deduction controls are ignored for designer contracts.",
                blocking=False,
            ))

    return issues


def validate_execution_subcontractor_controls(data: dict[str, Any]) -> list[ValidationIssue]:
    issues: list[ValidationIssue] = []
    c = data.get("subcontractor_controls", {})
    failure = c.get("failure_penalty", {})
    if not failure.get("amount") and not failure.get("percent"):
        issues.append(ValidationIssue("FAILURE_PENALTY_MISSING", "Failure penalty amount or percentage is mandatory."))
    guarantee = c.get("good_performance_guarantee", {})
    if not guarantee.get("amount") and not guarantee.get("percent"):
        issues.append(ValidationIssue("GOOD_PERFORMANCE_MISSING", "Good-performance guarantee amount or percentage is mandatory."))
    if Decimal(str(c.get("warranty_retention_percent", 0))) < Decimal("5"):
        issues.append(ValidationIssue("RETENTION_TOO_LOW", "Warranty retention must be at least 5% from every invoice."))
    if Decimal(str(c.get("client_services_deduction_percent", 0))) < Decimal("1"):
        issues.append(ValidationIssue("SERVICE_FEE_TOO_LOW", "Client-services deduction must be at least 1%."))
    if int(c.get("payment_term_days", 0) or 0) != EXECUTION_SUBCONTRACTOR_PAYMENT_TERM_DAYS:
        issues.append(ValidationIssue("EXECUTION_PAYMENT_TERM_INVALID", f"Execution subcontractor invoice payment term must be exactly {EXECUTION_SUBCONTRACTOR_PAYMENT_TERM_DAYS} calendar days."))
    if c.get("skonto_percent") is None:
        issues.append(ValidationIssue("SKONTO_UNDEFINED", "Possible skonto percentage must be explicitly defined, including 0%."))
    elif Decimal(str(c.get("skonto_percent"))) > 0 and int(c.get("skonto_days", 0) or 0) <= 0:
        issues.append(ValidationIssue("SKONTO_DAYS_MISSING", "A positive skonto requires an early-payment day limit."))
    if not _valid_https_url(str(c.get("subcontractor_terms_url", ""))):
        issues.append(ValidationIssue("ASZF_URL_INVALID", "A valid HTTPS subcontractor terms URL is mandatory."))
    gate = c.get("invoice_gate", {})
    for key, code, msg in [
        ("internal_tig_required", "INTERNAL_TIG_RULE_MISSING", "The signed and accepted internal TIG must be mandatory."),
        ("upstream_customer_tig_required", "UPSTREAM_TIG_RULE_MISSING", "The customer-signed upstream TIG must be mandatory."),
        ("risk_sharing_required", "RISK_SHARING_RULE_MISSING", "Contractual risk-sharing must be enabled."),
    ]:
        if gate.get(key) is not True:
            issues.append(ValidationIssue(code, msg))
    deductions = c.get("defect_deduction_schedule", [])
    if not deductions:
        issues.append(ValidationIssue("DEDUCTION_SCHEDULE_MISSING", "Technical-compliance deduction schedule cannot be empty."))
    for idx, item in enumerate(deductions, 1):
        if not item.get("defect"):
            issues.append(ValidationIssue("DEDUCTION_DEFECT_MISSING", f"Deduction row {idx} has no defect description."))
        if item.get("percent") is None and item.get("amount") is None:
            issues.append(ValidationIssue("DEDUCTION_VALUE_MISSING", f"Deduction row {idx} needs a concrete percentage or amount."))
    return issues


def assert_valid(data: dict[str, Any]) -> None:
    blocking = [i for i in validate_contract(data) if i.blocking]
    if blocking:
        raise ContractValidationError("; ".join(f"{i.code}: {i.message}" for i in blocking))


def dispatch_gate(data: dict[str, Any]) -> dict[str, Any]:
    blockers: list[str] = []
    status = data.get("dispatch_status", {})
    if status.get("internal_signed_original_present") is not True:
        blockers.append("INTERNALLY_SIGNED_ORIGINAL_MISSING")
    if not status.get("internal_signature_date") or not _valid_iso_date(status.get("internal_signature_date")):
        blockers.append("INTERNAL_SIGNATURE_DATE_MISSING_OR_INVALID")
    if not status.get("signed_document_sha256") or not re.fullmatch(r"[0-9a-fA-F]{64}", str(status.get("signed_document_sha256", ""))):
        blockers.append("SIGNED_DOCUMENT_HASH_MISSING_OR_INVALID")

    postal = status.get("postal", {})
    if postal.get("sent") is not True:
        blockers.append("SIGNED_ORIGINAL_NOT_POSTED")
    if not postal.get("sent_at"):
        blockers.append("POSTAL_SENT_AT_MISSING")
    if not postal.get("tracking_number"):
        blockers.append("POSTAL_TRACKING_NUMBER_MISSING")
    if not postal.get("proof_file_id"):
        blockers.append("POSTAL_PROOF_MISSING")
    if _normalize(postal.get("recipient_address")) != _normalize(data.get("counterparty", {}).get("postal_address")):
        blockers.append("POSTAL_RECIPIENT_ADDRESS_MISMATCH")
    if int(postal.get("original_copy_count", 0) or 0) < int(data.get("delivery_requirements", {}).get("postal_original_copies", 1) or 1):
        blockers.append("POSTAL_ORIGINAL_COPY_COUNT_INSUFFICIENT")

    electronic = status.get("electronic", {})
    if electronic.get("sent") is not True:
        blockers.append("SIGNED_CONTRACT_NOT_SENT_ELECTRONICALLY")
    if not electronic.get("sent_at"):
        blockers.append("ELECTRONIC_SENT_AT_MISSING")
    if _normalize(electronic.get("recipient_email")) != _normalize(data.get("counterparty", {}).get("email")):
        blockers.append("ELECTRONIC_RECIPIENT_MISMATCH")
    if not electronic.get("message_id"):
        blockers.append("ELECTRONIC_MESSAGE_ID_MISSING")
    if electronic.get("attachment_sha256") != status.get("signed_document_sha256"):
        blockers.append("ELECTRONIC_ATTACHMENT_HASH_MISMATCH")
    return {"allowed": not blockers, "blockers": list(dict.fromkeys(blockers))}

def work_start_gate(data: dict[str, Any]) -> dict[str, Any]:
    blockers: list[str] = []
    status = data.get("status", {})
    if status.get("contract_status") != "SIGNED":
        blockers.append("CONTRACT_NOT_SIGNED")
    if status.get("signed_contract_present") is not True:
        blockers.append("SIGNED_FILE_MISSING")
    if status.get("both_parties_signed") is not True:
        blockers.append("BOTH_PARTIES_SIGNATURE_MISSING")
    if not status.get("signed_contract_file_id"):
        blockers.append("SIGNED_CONTRACT_FILE_ID_MISSING")
    if status.get("master_hash_verified") is not True:
        blockers.append("MASTER_INTEGRITY_FAILED")
    if status.get("all_required_annexes_present") is not True:
        blockers.append("MANDATORY_ANNEX_MISSING")
    if status.get("commercial_approval") != "APPROVED":
        blockers.append("COMMERCIAL_APPROVAL_MISSING")
    if status.get("technical_approval") != "APPROVED":
        blockers.append("TECHNICAL_APPROVAL_MISSING")
    if status.get("all_fields_complete") is not True:
        blockers.append("CONTRACT_FIELDS_INCOMPLETE")
    dispatch = dispatch_gate(data)
    blockers.extend(dispatch["blockers"])
    return {"allowed": not blockers, "blockers": list(dict.fromkeys(blockers)), "dispatch_gate": dispatch}


def _invoice_required_fields() -> list[str]:
    return [
        "invoice_number", "invoice_type", "issue_date", "performance_date", "received_date", "due_date",
        "currency", "payment_method", "supplier.name", "supplier.address", "supplier.tax_number",
        "supplier.bank_account", "buyer.name", "buyer.address", "buyer.tax_number",
        "net_amount", "vat_percent", "vat_amount", "gross_amount",
        "contract_number", "project_id", "performance_certificate.id",
        "performance_certificate.contract_number", "performance_certificate.project_id",
        "performance_certificate.status", "performance_certificate.accepted_date",
        "performance_certificate.accepted_by_name", "performance_certificate.accepted_by_title",
        "performance_certificate.supplier_signed", "performance_certificate.authorized_acceptor_signed",
        "performance_certificate.accepted_gross_amount", "line_items", "file_present", "file_sha256",
        "document_readable", "content_integrity_verified",
    ]

def _normalize(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().casefold()


def invoice_acceptance_gate(data: dict[str, Any], invoice: dict[str, Any]) -> dict[str, Any]:
    if data.get("contract_type") not in SUBCONTRACTOR_TYPES:
        raise ContractValidationError("Invoice acceptance gate is available only for subcontractor/partner contracts.")

    blockers: list[str] = []
    details: dict[str, str] = {}
    for field in _missing(invoice, _invoice_required_fields()):
        blockers.append("INVOICE_FIELD_MISSING")
        details[field] = "Kötelező számlaadat vagy igazolási adat hiányzik."

    for field in ("issue_date", "performance_date", "received_date", "due_date"):
        if invoice.get(field) and not _valid_iso_date(invoice.get(field)):
            blockers.append("INVOICE_DATE_INVALID")
            details[field] = "A dátum nem érvényes ISO dátum."

    if invoice.get("invoice_type") not in {"INVOICE", "ADVANCE_INVOICE", "FINAL_INVOICE", "CORRECTION_INVOICE"}:
        blockers.append("INVOICE_TYPE_INVALID")
    if invoice.get("currency") != data.get("commercial", {}).get("currency"):
        blockers.append("INVOICE_CURRENCY_MISMATCH")
    if invoice.get("file_present") is not True:
        blockers.append("INVOICE_FILE_MISSING")
    if not re.fullmatch(r"[0-9a-fA-F]{64}", str(invoice.get("file_sha256", ""))):
        blockers.append("INVOICE_FILE_HASH_INVALID")
    if invoice.get("document_readable") is not True:
        blockers.append("INVOICE_NOT_READABLE")
    if invoice.get("content_integrity_verified") is not True:
        blockers.append("INVOICE_INTEGRITY_NOT_VERIFIED")
    if invoice.get("duplicate_invoice") is True:
        blockers.append("DUPLICATE_INVOICE")

    cp = data.get("counterparty", {})
    internal = data.get("internal_entity", {})
    supplier = invoice.get("supplier", {})
    buyer = invoice.get("buyer", {})
    comparisons = [
        ("supplier.name", supplier.get("name"), cp.get("name"), "SUPPLIER_NAME_MISMATCH"),
        ("supplier.address", supplier.get("address"), cp.get("address"), "SUPPLIER_ADDRESS_MISMATCH"),
        ("supplier.tax_number", supplier.get("tax_number"), cp.get("tax_number"), "SUPPLIER_TAX_MISMATCH"),
        ("supplier.bank_account", supplier.get("bank_account"), cp.get("bank_account"), "SUPPLIER_BANK_ACCOUNT_MISMATCH"),
        ("buyer.name", buyer.get("name"), internal.get("name"), "BUYER_NAME_MISMATCH"),
        ("buyer.address", buyer.get("address"), internal.get("address"), "BUYER_ADDRESS_MISMATCH"),
        ("buyer.tax_number", buyer.get("tax_number"), internal.get("tax_number"), "BUYER_TAX_MISMATCH"),
    ]
    for field, actual, expected, code in comparisons:
        if actual and expected and _normalize(actual) != _normalize(expected):
            blockers.append(code)
            details[field] = "A számlán szereplő adat nem egyezik a szerződéses törzsadattal."

    if invoice.get("contract_number") != data.get("contract_number"):
        blockers.append("CONTRACT_REFERENCE_MISMATCH")
    if invoice.get("project_id") != data.get("ids", {}).get("ProjectID"):
        blockers.append("PROJECT_REFERENCE_MISMATCH")

    pc = invoice.get("performance_certificate", {})
    if pc.get("contract_number") != data.get("contract_number"):
        blockers.append("TIG_CONTRACT_REFERENCE_MISMATCH")
    if pc.get("project_id") != data.get("ids", {}).get("ProjectID"):
        blockers.append("TIG_PROJECT_REFERENCE_MISMATCH")
    if pc.get("status") != "ACCEPTED":
        blockers.append("TIG_NOT_ACCEPTED")
    if pc.get("supplier_signed") is not True:
        blockers.append("TIG_SUPPLIER_SIGNATURE_MISSING")
    if pc.get("authorized_acceptor_signed") is not True:
        blockers.append("TIG_AUTHORIZED_SIGNATURE_MISSING")
    if not pc.get("accepted_by_name") or not pc.get("accepted_by_title"):
        blockers.append("TIG_AUTHORIZED_ACCEPTOR_DATA_MISSING")

    try:
        net = money(invoice.get("net_amount", 0))
        vat_pct = Decimal(str(invoice.get("vat_percent", 0)))
        vat = money(invoice.get("vat_amount", 0))
        gross = money(invoice.get("gross_amount", 0))
        if net <= 0 or gross <= 0:
            blockers.append("INVOICE_AMOUNT_INVALID")
        if abs(vat - pct_amount(net, vat_pct)) > Decimal("1.00"):
            blockers.append("INVOICE_VAT_MISMATCH")
        if abs(gross - (net + vat)) > Decimal("1.00"):
            blockers.append("INVOICE_TOTAL_MISMATCH")
        if pc.get("accepted_gross_amount") is not None and abs(gross - money(pc.get("accepted_gross_amount"))) > Decimal("1.00"):
            blockers.append("INVOICE_EXCEEDS_OR_DIFFERS_FROM_TIG")
    except Exception:
        blockers.append("INVOICE_AMOUNT_FORMAT_INVALID")
        net = vat = gross = Decimal("0")
        vat_pct = Decimal("0")

    line_net = line_vat = line_gross = Decimal("0")
    if not invoice.get("line_items"):
        blockers.append("INVOICE_LINE_ITEMS_MISSING")
    else:
        for idx, item in enumerate(invoice.get("line_items", []), 1):
            required_line = ("description", "quantity", "unit", "unit_net_price", "net_amount", "vat_percent", "vat_amount", "gross_amount")
            if any(item.get(field) is None or (isinstance(item.get(field), str) and not item.get(field).strip()) for field in required_line):
                blockers.append("INVOICE_LINE_ITEM_INVALID")
                details[f"line_items.{idx}"] = "A tételsor megnevezése, mennyisége, egysége, egységára vagy adóadata hiányzik."
                continue
            try:
                qty = Decimal(str(item["quantity"]))
                unit_net = money(item["unit_net_price"])
                row_net = money(item["net_amount"])
                row_vat_pct = Decimal(str(item["vat_percent"]))
                row_vat = money(item["vat_amount"])
                row_gross = money(item["gross_amount"])
                if abs(row_net - (qty * unit_net).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)) > Decimal("1.00"):
                    blockers.append("INVOICE_LINE_NET_MISMATCH")
                if abs(row_vat - pct_amount(row_net, row_vat_pct)) > Decimal("1.00"):
                    blockers.append("INVOICE_LINE_VAT_MISMATCH")
                if abs(row_gross - (row_net + row_vat)) > Decimal("1.00"):
                    blockers.append("INVOICE_LINE_GROSS_MISMATCH")
                line_net += row_net; line_vat += row_vat; line_gross += row_gross
            except Exception:
                blockers.append("INVOICE_LINE_ITEM_FORMAT_INVALID")
        if abs(line_net - net) > Decimal("1.00"):
            blockers.append("INVOICE_LINE_SUM_NET_MISMATCH")
        if abs(line_vat - vat) > Decimal("1.00"):
            blockers.append("INVOICE_LINE_SUM_VAT_MISMATCH")
        if abs(line_gross - gross) > Decimal("1.00"):
            blockers.append("INVOICE_LINE_SUM_GROSS_MISMATCH")

    if data.get("contract_type") == EXECUTION_SUBCONTRACTOR_TYPE:
        c = data.get("subcontractor_controls", {})
        gate = c.get("invoice_gate", {})
        if gate.get("upstream_customer_tig_required") and invoice.get("upstream_customer_tig_signed") is not True:
            blockers.append("UPSTREAM_CUSTOMER_TIG_NOT_SIGNED")
        if gate.get("risk_sharing_required"):
            if invoice.get("upstream_invoice_accepted") is not True:
                blockers.append("UPSTREAM_INVOICE_NOT_ACCEPTED")
            if gate.get("upstream_payment_required", True) and invoice.get("upstream_invoice_paid") is not True:
                blockers.append("UPSTREAM_INVOICE_NOT_PAID")
        payment_days = int(c.get("payment_term_days", 0) or 0)
        if payment_days != EXECUTION_SUBCONTRACTOR_PAYMENT_TERM_DAYS:
            blockers.append("EXECUTION_PAYMENT_TERM_INVALID")
        retention = pct_amount(gross, c.get("warranty_retention_percent", 0))
        service_fee = pct_amount(gross, c.get("client_services_deduction_percent", 0))
        guarantee_spec = c.get("good_performance_guarantee", {})
        good_performance = pct_amount(gross, guarantee_spec.get("percent")) if guarantee_spec.get("percent") is not None else money(guarantee_spec.get("amount", 0))
        payable = gross - retention - service_fee - good_performance
        skonto_percent = Decimal(str(c.get("skonto_percent", 0) or 0))
        skonto_amount = pct_amount(payable, skonto_percent) if skonto_percent > 0 else Decimal("0")
    else:
        payment_days = int(data.get("invoice_controls", {}).get("payment_term_days", 0) or 0)
        if data.get("contract_type") == DESIGN_SUBCONTRACTOR_TYPE and payment_days != DESIGNER_PAYMENT_TERM_DAYS:
            blockers.append("DESIGN_PAYMENT_TERM_INVALID")
        elif payment_days <= 0:
            blockers.append("PAYMENT_TERM_INVALID")
        retention = service_fee = good_performance = skonto_amount = Decimal("0")
        payable = gross
        skonto_percent = Decimal("0")

    base_date_value = invoice.get("received_date")
    base_date = date.fromisoformat(base_date_value) if base_date_value and _valid_iso_date(base_date_value) else date.today()
    standard_due_date = base_date + timedelta(days=payment_days)
    if invoice.get("due_date") and _valid_iso_date(invoice.get("due_date")):
        stated_due = date.fromisoformat(invoice["due_date"])
        if stated_due < standard_due_date:
            blockers.append("INVOICE_DUE_DATE_TOO_EARLY")
            details["due_date"] = f"A szerződés szerinti legkorábbi esedékesség: {standard_due_date.isoformat()}."

    result = {
        "accepted": not blockers,
        "status": "ACCEPTED" if not blockers else "REJECTED_IMMEDIATELY",
        "blockers": list(dict.fromkeys(blockers)),
        "field_details": details,
        "calculation": {
            "invoice_gross": str(gross.quantize(Decimal("0.01"))),
            "warranty_retention": str(retention.quantize(Decimal("0.01"))),
            "client_services_deduction": str(service_fee.quantize(Decimal("0.01"))),
            "good_performance_guarantee": str(good_performance.quantize(Decimal("0.01"))),
            "standard_payable": str(payable.quantize(Decimal("0.01"))),
            "standard_due_date": standard_due_date.isoformat(),
            "skonto_percent": str(skonto_percent),
            "skonto_amount": str(skonto_amount.quantize(Decimal("0.01"))),
        },
    }
    if blockers:
        result["rejection_notice"] = build_invoice_rejection_notice(data, invoice, result)
    return result

def build_invoice_rejection_notice(data: dict[str, Any], invoice: dict[str, Any], gate_result: dict[str, Any]) -> dict[str, Any]:
    blocker_labels = {
        "INVOICE_FIELD_MISSING": "a számla kötelező vagy szerződésben előírt adatai hiányosak",
        "INVOICE_NOT_READABLE": "a számla nem olvasható",
        "INVOICE_INTEGRITY_NOT_VERIFIED": "a számla adattartalmának sértetlensége nem igazolt",
        "TIG_NOT_ACCEPTED": "a teljesítésigazolás nincs elfogadott állapotban",
        "TIG_SUPPLIER_SIGNATURE_MISSING": "a teljesítésigazolás partneri aláírása hiányzik",
        "TIG_AUTHORIZED_SIGNATURE_MISSING": "a teljesítésigazolás megrendelői/jogosult aláírása hiányzik",
        "INVOICE_EXCEEDS_OR_DIFFERS_FROM_TIG": "a számla összege nem egyezik az elfogadott teljesítésigazolással",
        "UPSTREAM_CUSTOMER_TIG_NOT_SIGNED": "a szerződésben előírt megrendelői upstream TIG hiányzik",
        "DUPLICATE_INVOICE": "a számla sorszáma már szerepel a rendszerben",
        "SUPPLIER_NAME_MISMATCH": "a szállító neve nem egyezik a szerződéses törzsadattal",
        "SUPPLIER_ADDRESS_MISMATCH": "a szállító címe nem egyezik a szerződéses törzsadattal",
        "SUPPLIER_TAX_MISMATCH": "a szállító adószáma nem egyezik a szerződéses törzsadattal",
        "SUPPLIER_BANK_ACCOUNT_MISMATCH": "a szállító bankszámlaszáma nem egyezik a szerződéses törzsadattal",
        "BUYER_NAME_MISMATCH": "a vevő neve nem egyezik a szerződéses törzsadattal",
        "BUYER_ADDRESS_MISMATCH": "a vevő címe nem egyezik a szerződéses törzsadattal",
        "BUYER_TAX_MISMATCH": "a vevő adószáma nem egyezik a szerződéses törzsadattal",
        "CONTRACT_REFERENCE_MISMATCH": "a számla szerződéshivatkozása hibás",
        "PROJECT_REFERENCE_MISMATCH": "a számla projekthivatkozása hibás",
        "TIG_CONTRACT_REFERENCE_MISMATCH": "a teljesítésigazolás szerződéshivatkozása hibás",
        "TIG_PROJECT_REFERENCE_MISMATCH": "a teljesítésigazolás projekthivatkozása hibás",
        "INVOICE_DUE_DATE_TOO_EARLY": "a számlán feltüntetett esedékesség korábbi a szerződés szerint megengedettnél",
        "INVOICE_LINE_ITEM_INVALID": "a számla tételsora hiányos",
        "INVOICE_LINE_SUM_NET_MISMATCH": "a tételsorok nettó összege nem egyezik a számla nettó végösszegével",
        "INVOICE_LINE_SUM_VAT_MISMATCH": "a tételsorok áfaösszege nem egyezik a számla áfa-végösszegével",
        "INVOICE_LINE_SUM_GROSS_MISMATCH": "a tételsorok bruttó összege nem egyezik a számla bruttó végösszegével",
        "INVOICE_VAT_MISMATCH": "a számla áfaösszege számtanilag hibás",
        "INVOICE_TOTAL_MISMATCH": "a számla bruttó végösszege számtanilag hibás",
        "UPSTREAM_INVOICE_NOT_ACCEPTED": "a szerződés szerinti upstream számlát a megrendelő még nem fogadta be",
        "UPSTREAM_INVOICE_NOT_PAID": "a szerződés szerinti upstream ellenérték még nem érkezett meg",
    }
    reasons = [blocker_labels.get(code, code.replace("_", " ").lower()) for code in gate_result.get("blockers", [])]
    reason_text = "; ".join(reasons)
    subject = f"Számla befogadásának elutasítása – {invoice.get('invoice_number', 'azonosítatlan számla')}"
    body = (
        f"Tisztelt {data.get('counterparty', {}).get('name', 'Partner')}!\n\n"
        f"Tájékoztatjuk, hogy a(z) {invoice.get('invoice_number', 'azonosítatlan')} sorszámú számlát nem fogadjuk be, "
        "azt könyvelésre és kifizetésre nem továbbítjuk. Az elutasítás oka: " + reason_text + ".\n\n"
        "A szerződés szerint számla kizárólag a szerződésszerű teljesítést igazoló, a jogosult személyek által aláírt és elfogadott "
        "teljesítésigazolás alapján, a teljesítésigazolással egyező tartalommal és összeggel nyújtható be. "
        "A számlának meg kell felelnie az Áfa tv. 168/A. §-ában foglalt hitelességi, sértetlenségi és olvashatósági követelményeknek, "
        "valamint az Áfa tv. 169. §-ában meghatározott kötelező adattartalomnak.\n\n"
        "Kérjük, a hibás számla tekintetében a szükséges sztornó- vagy módosító bizonylatot állítsák ki, illetve a hiányzó, aláírt és elfogadott "
        "teljesítésigazolással együtt a javított számlát új benyújtásként küldjék meg. A számla módosítására szolgáló okiratnak az Áfa tv. 170. §-a "
        "szerinti adatokat kell tartalmaznia. A fizetési határidő kizárólag a hibátlan, hiánytalan és befogadott számla rendszerben rögzített "
        "befogadási napjától kezdődhet.\n\n"
        "Jelen értesítés nem minősül a teljesítés vagy a követelés elismerésének.\n\n"
        "Üdvözlettel:\nImperial Holding"
    )
    return {
        "subject": subject,
        "body": body,
        "legal_basis": LEGAL_SOURCE_NOTES,
        "required_action": "Sztornó/módosító bizonylat és hibátlan új benyújtás; hiányzó aláírt, elfogadott TIG csatolása.",
    }


def _set_paragraph(doc: DocxDocument, index: int, text: str, audit: list[dict[str, Any]], field: str) -> None:
    if index < len(doc.paragraphs):
        paragraph = doc.paragraphs[index]
        paragraph.text = text
        for run in paragraph.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
        audit.append({"field": field, "changed": True, "target": f"paragraph:{index}"})
    else:
        audit.append({"field": field, "changed": False, "target": f"paragraph:{index}"})


def _set_table_value(doc: DocxDocument, table_index: int, row_index: int, value: str, audit: list[dict[str, Any]], field: str) -> None:
    try:
        row = doc.tables[table_index].rows[row_index]
        target = 1 if len(row.cells) == 2 else 2
        row.cells[target].text = str(value)
        audit.append({"field": field, "changed": True, "target": f"table:{table_index}:row:{row_index}:cell:{target}"})
    except Exception:
        audit.append({"field": field, "changed": False, "target": f"table:{table_index}:row:{row_index}"})


def _format_hu_date(value: str, place: str = "Budapest") -> str:
    d = date.fromisoformat(value)
    return f"{place}, {d.year}. {d.month:02d}. {d.day:02d}."


def _party_description(p: dict[str, Any]) -> str:
    if p.get("party_type") == "natural_person":
        return (
            f"{p['name']} (születési hely, idő: {p['birth_place']}, {p['birth_date']}; anyja neve: {p['mother_name']}; "
            f"{p['identity_document_type']} száma: {p['identity_document_number']}; lakcím: {p['address']}; "
            f"e-mail: {p['email']}; telefon: {p['phone']})"
        )
    return (
        f"{p['name']} (rövid név: {p['short_name']}; székhely: {p['registered_office']}; cégjegyzékszám/nyilvántartási szám: "
        f"{p['registration_number']}; adószám: {p['tax_number']}; bankszámlaszám: {p['bank_account']}; "
        f"képviseli: {p['representative']} {p['representative_title']}; e-mail: {p['email']}; telefon: {p['phone']})"
    )


ATTACHMENT_LABELS = {
    "technical_scope": "műszaki tartalom", "priced_bill_of_quantities": "árazott költségvetés",
    "priced_budget": "árazott költségvetés", "schedule": "ütemterv", "plans": "tervdokumentáció",
    "technical_compliance_deductions": "műszaki megfelelőségi és levonási melléklet",
    "invoice_gate_terms": "számlabefogadási feltételek", "design_scope": "tervezési program és feladatleírás",
    "deliverables_schedule": "tervszállítási és mérföldkő-ütemterv", "fee_schedule": "díjazási és fizetési ütem",
    "professional_liability_insurance": "tervezői felelősségbiztosítás igazolása",
    "design_program": "tervezési program", "deliverables": "átadandó tervdokumentációk jegyzéke",
    "type_house_plan": "típusház tervdokumentáció",
}


def _annex_names(data: dict[str, Any]) -> str:
    return "; ".join(
        f"{ATTACHMENT_LABELS.get(x.get('type'), x.get('type'))} (v{x.get('version', '-')}, {x.get('status')})"
        for x in data.get("attachments", [])
    )


def _payment_schedule_sentence(data: dict[str, Any]) -> str:
    parts = []
    for row in data.get("commercial", {}).get("payment_schedule", []):
        parts.append(f"{row['milestone']}: {row['percent']}%, esedékesség: {row['due_rule']}")
    return "; ".join(parts)


def _clear_designer_review_markup(doc: DocxDocument) -> None:
    """Remove template-review coloring/highlighting from the generated designer copy."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = None
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                run.font.color.rgb = RGBColor(0, 0, 0)


def fill_designated_fields(template_path: Path, output_path: Path, data: dict[str, Any]) -> list[dict[str, Any]]:
    shutil.copy2(template_path, output_path)
    doc = create_docx_document(str(output_path))
    audit: list[dict[str, Any]] = []
    ct = data["contract_type"]
    cp, ie, pr, c, sch = data["counterparty"], data["internal_entity"], data["project"], data["commercial"], data["schedule"]
    site = f"{pr['site_address']}, {pr['parcel_number']} hrsz."
    price = f"nettó {money(c['net_price']):,.0f} {c['currency']} + {c['vat_percent']}% ÁFA = bruttó {money(c['gross_price']):,.0f} {c['currency']}".replace(",", " ")

    if ct == "customer_type_house_design_build":
        values = [
            cp["name"], f"{cp['birth_place']}, {cp['birth_date']}", f"{cp['identity_document_type']} {cp['identity_document_number']}",
            cp["mother_name"], cp["phone"], cp["email"], cp["address"], cp["name"], cp["phone"],
            ie["name"], ie["registered_office"], ie["tax_number"], ie["representative"], ie["phone"], ie["email"],
            site, sch["deadline"], price, sch.get("site_handover_date", sch["start_date"]), _annex_names(data),
        ]
        for i, value in enumerate(values):
            _set_table_value(doc, 0, i, value, audit, f"table_row_{i}")
        _set_paragraph(doc, 143, _format_hu_date(data["contract_date"], data["contract_place"]), audit, "contract_date")
        _set_paragraph(doc, 147, f"{cp['name']}                                      {ie['representative']} / {ie['name']}", audit, "signature_names")
        _set_paragraph(doc, 148, "", audit, "remove_stale_signer")
        _set_paragraph(doc, 149, "", audit, "remove_stale_signature_company")

    elif ct == "customer_construction":
        _set_paragraph(doc, 4, f"Név: {cp['name']}", audit, "counterparty.name")
        _set_paragraph(doc, 5, f"Szül. hely, idő: {cp['birth_place']}, {cp['birth_date']}", audit, "counterparty.birth")
        _set_paragraph(doc, 6, f"Anyja neve: {cp['mother_name']}", audit, "counterparty.mother_name")
        _set_paragraph(doc, 7, f"Lakcím: {cp['address']}", audit, "counterparty.address")
        _set_paragraph(doc, 8, f"{cp['identity_document_type']} száma: {cp['identity_document_number']}", audit, "counterparty.id")
        _set_paragraph(doc, 9, f"E-mail: {cp['email']}; telefon: {cp['phone']}", audit, "counterparty.contact")
        _set_paragraph(doc, 13, f"cégnév: {ie['name']}", audit, "internal.name")
        _set_paragraph(doc, 14, f"Székhely: {ie['registered_office']}", audit, "internal.address")
        _set_paragraph(doc, 15, f"Adószám: {ie['tax_number']}; cégjegyzékszám: {ie['registration_number']}", audit, "internal.registry")
        _set_paragraph(doc, 16, f"Képviseli: {ie['representative']} {ie['representative_title']}", audit, "internal.representative")
        _set_paragraph(doc, 17, f"E-mail: {ie['email']}; telefon: {ie['phone']}", audit, "internal.contact")
        _set_paragraph(doc, 18, f"Bankszámlaszám: {ie['bank_account']}", audit, "internal.bank")
        table_values = [site, sch["deadline"], price, sch["site_handover_date"], cp["name"], cp["phone"], ie["representative"], ie["phone"], ie["email"], _annex_names(data)]
        for i, value in enumerate(table_values):
            _set_table_value(doc, 0, i, value, audit, f"table_row_{i}")
        _set_paragraph(doc, 148, _format_hu_date(data["contract_date"], data["contract_place"]), audit, "contract_date")
        _set_paragraph(doc, 154, f"{cp['name']}                                      {ie['representative']} / {ie['name']}", audit, "signature_names")
        _set_paragraph(doc, 155, "", audit, "remove_stale_signature_company")

    elif ct == "customer_design_execution_plans":
        _set_paragraph(doc, 1, f"amely létrejött egyrészről {_party_description(cp)}, mint Megrendelő,", audit, "counterparty")
        _set_paragraph(doc, 2, f"másrészről {_party_description(ie)}, mint Vállalkozó", audit, "internal_entity")
        _set_paragraph(doc, 4, f"1) A Megrendelő megrendeli és Vállalkozó elvállalja a {pr['parcel_number']} helyrajzi szám alatt, természetben a {pr['site_address']} cím alatt található, bruttó {pr['gross_floor_area_m2']} m² alapterületű lakóház {pr['scope']} tárgyú építési terveinek elkészítését.", audit, "project")
        _set_paragraph(doc, 17, f"4) A tervezés díja egy összegben: {price}.", audit, "price")
        original = doc.paragraphs[31].text
        original = re.sub(r"_{5,}-ig", f"{sch['deadline']}-ig", original)
        _set_paragraph(doc, 31, original, audit, "deadline")
        payment_text = (
            f"16) Fizetési feltételek: A teljes tervezési díj {price}. A fizetés kizárólag a Vállalkozó "
            f"{ie['bank_account']} számú bankszámlájára történő átutalással teljesíthető. Fizetési ütem: "
            f"{_payment_schedule_sentence(data)}. A Vállalkozó az adott tervfázis dokumentációját a kapcsolódó "
            "fizetési kötelezettség teljesítéséig visszatarthatja."
        )
        _set_paragraph(doc, 36, payment_text, audit, "payment_schedule_and_bank_account")
        _set_paragraph(doc, 48, _format_hu_date(data["contract_date"], data["contract_place"]), audit, "contract_date")
        _set_paragraph(doc, 51, f"{cp['name']}                                      {ie['representative']} / {ie['name']}", audit, "signature_names")
        _set_paragraph(doc, 52, "", audit, "remove_stale_signer")
        _set_paragraph(doc, 53, "", audit, "remove_stale_signature_company")

    elif ct == EXECUTION_SUBCONTRACTOR_TYPE:
        rows = [
            (6, f"cégnév: {ie['name']}"), (7, f"rövid cégnév: {ie['short_name']}"),
            (8, f"cégjegyzékszám: {ie['registration_number']}"), (9, f"adószám: {ie['tax_number']}"),
            (10, f"székhely: {ie['registered_office']}"), (11, f"képviseli: {ie['representative']} {ie['representative_title']}"),
            (16, f"cégnév: {cp['name']}"), (17, f"rövid cégnév: {cp['short_name']}"),
            (18, f"cégjegyzékszám: {cp['registration_number']}"), (19, f"adószám: {cp['tax_number']}"),
            (20, f"székhely: {cp['registered_office']}"), (21, f"képviseli: {cp['representative']} {cp['representative_title']}"),
        ]
        for idx, text in rows:
            _set_paragraph(doc, idx, text, audit, text.split(":")[0])
        _set_paragraph(doc, 29, f"{pr['scope']} elvégzése a {site} alatti projekten.", audit, "project.scope")
        original = doc.paragraphs[86].text
        original = re.sub(r"Alvállalkozót … vállalkozói díj illeti meg", f"Alvállalkozót {price} vállalkozói díj illeti meg", original)
        _set_paragraph(doc, 86, original, audit, "price")
        _set_paragraph(doc, 115, f"5.1. A munkaterület átadása: {sch['site_handover_date']}", audit, "site_handover_date")
        _set_paragraph(doc, 116, f"5.2. A munkakezdés időpontja: {sch['start_date']}", audit, "start_date")
        _set_paragraph(doc, 117, f"5.3. Befejezési véghatáridő/műszaki átadás-átvétel: {sch['deadline']}", audit, "deadline")
        _set_paragraph(doc, 59, doc.paragraphs[59].text.replace("info@imperialholding.hu", ie["email"]), audit, "internal_operational_email")
        _set_paragraph(doc, 124, f"Megrendelő részéről: {ie['representative']} {ie['representative_title']}, e-mail: {ie['email']}, telefon: {ie['phone']}", audit, "internal_contact")
        _set_paragraph(doc, 126, f"Alvállalkozó részéről: {cp['representative']}, e-mail: {cp['email']}, telefon: {cp['phone']}", audit, "partner_contact")
        _set_paragraph(doc, 160, doc.paragraphs[160].text.replace("info@imperialholding.hu", ie["email"]), audit, "internal_notice_email")
        _set_paragraph(doc, 168, _format_hu_date(data["contract_date"], data["contract_place"]), audit, "contract_date")
        _set_paragraph(doc, 171, f"{ie['representative']} / {ie['name']}                         {cp['representative']} / {cp['name']}", audit, "signature_names")
        _set_paragraph(doc, 174, "", audit, "remove_witness_intro")
        _set_paragraph(doc, 175, "Tanúk: nem alkalmazandók – mindkét fél cégszerű aláírással jár el.", audit, "witness_heading")
        for idx in (176, 177, 178, 179):
            _set_paragraph(doc, idx, "", audit, f"witness_{idx}")

    elif ct == DESIGN_SUBCONTRACTOR_TYPE:
        _set_paragraph(doc, 4, f"amely létrejött egyrészről {_party_description(ie)}, mint Megrendelő,", audit, "internal_entity")
        _set_paragraph(doc, 6, f"másrészről {_party_description(cp)}, mint Vállalkozó – a továbbiakban: Vállalkozó", audit, "counterparty")
        _set_paragraph(doc, 12, f"A Megrendelő megrendeli, és Vállalkozó elvállalja a {pr['parcel_number']} helyrajzi számon nyilvántartott, természetben a {pr['site_address']} cím alatti, {pr['gross_floor_area_m2']} m² alapterületű ingatlanon a következő tervezési feladatot: {pr['scope']}. Az eljárás típusa: {pr['procedure_type']}.", audit, "project")
        procedure = str(pr["procedure_type"]).strip().casefold()
        procedure_label = "építési engedélyezési eljárás" if "engedély" in procedure else "egyszerű bejelentési eljárás"
        _set_paragraph(doc, 16, f"A jelen szerződésben részletezett szolgáltatások az épület, épületrész, épületegyüttes és a hozzá tartozó építmények megépítéséhez szükséges, jogszabályban meghatározott tartalmú műszaki tervdokumentáció elkészítését tartalmazzák a {procedure_label} követelményeire vonatkozóan.", audit, "procedure_clause")
        _set_paragraph(doc, 29, f"A tervezés díja: {price}, amelyre a Vállalkozó a szerződésszerű teljesítés ellenértékeként jogosult.", audit, "price")
        ins = data.get("designer_controls", {}).get("professional_liability_insurance", {})
        _set_paragraph(doc, 52, f"Vállalkozó szakmai felelősségbiztosítója: {ins.get('insurer')}; kötvényszám: {ins.get('policy_number')}; fedezeti értékhatár: {ins.get('coverage_amount')} {c['currency']}. A biztosítás igazolása a szerződés kötelező melléklete.", audit, "insurance")
        milestones = sch.get("milestones", [])
        for offset in range(3):
            row = milestones[offset] if offset < len(milestones) else {"name": "Nem alkalmazandó", "deadline": "Nem alkalmazandó"}
            _set_paragraph(doc, 95 + offset, f"{offset+1}. {row.get('name')}: {row.get('deadline')}", audit, f"milestone_{offset+1}")
        _set_paragraph(doc, 99, f"Az 1. pontban meghatározott tervezési feladatok elvégzésének teljesítési határideje: {sch['deadline']}.", audit, "deadline")
        _set_paragraph(doc, 150, f"Megrendelő e-mail címe: {ie['email']}, Vállalkozó e-mail címe: {cp['email']}", audit, "emails")
        _set_paragraph(doc, 158, f"Kelt: {_format_hu_date(data['contract_date'], data['contract_place'])}", audit, "contract_date")
        _set_paragraph(doc, 161, f"{ie['representative']} / {ie['name']}                         {cp['representative']} / {cp['name']}", audit, "signature_names")
        if ie.get("party_type") in {"company", "sole_proprietor"} and cp.get("party_type") in {"company", "sole_proprietor"}:
            _set_paragraph(doc, 169, "Tanúk: nem alkalmazandók – a felek cégszerű aláírással járnak el.", audit, "witness_heading")
            for idx in range(170, 176):
                _set_paragraph(doc, idx, "", audit, f"witness_{idx}")
        _clear_designer_review_markup(doc)

    # Mandatory generated clauses are inserted into the generated copy, never into the immutable master.
    date_index_by_type = {
        "customer_type_house_design_build": 143,
        "customer_construction": 148,
        "customer_design_execution_plans": 48,
        "subcontractor_execution": 168,
        "subcontractor_design": 158,
    }
    insertion_index = date_index_by_type.get(ct, len(doc.paragraphs) - 1)
    target = doc.paragraphs[insertion_index]
    delivery_clause = (
        "KÖTELEZŐ KÉZBESÍTÉSI ZÁRADÉK: Az Imperial által cégszerűen aláírt szerződés legalább egy eredeti példányát "
        "igazolható, nyomkövethető postai vagy futárküldeményként a partner hivatalos postai címére, "
        "és azzal egyidejűleg az aláírt elektronikus példányt a partner hivatalos e-mail-címére is meg kell küldeni. "
        "A nyomkövetési számot, feladási vagy kézbesítési bizonylatot, az elektronikus üzenet azonosítóját, "
        "a címzettet, a küldési időt és a csatolmány SHA-256 értékét a szerződésnyilvántartásban rögzíteni kell. "
        "A kettős kézbesítés igazolásáig munkakezdési engedély nem adható."
    )
    target.insert_paragraph_before(delivery_clause)
    audit.append({"field": "mandatory_delivery_clause", "changed": True, "target": f"before_paragraph:{insertion_index}"})
    if ct in SUBCONTRACTOR_TYPES:
        invoice_clause = (
            "Kötelező számlabefogadási záradék: Számla kizárólag a jogosult személyek által aláírt és elfogadott "
            "teljesítésigazolás alapján, azzal egyező tartalommal és összeggel, hibátlan formai és tartalmi adatokkal nyújtható be. "
            "Hiányos, hibás, olvashatatlan, sérült, duplikált, a szerződéses törzsadatoktól vagy a teljesítésigazolástól eltérő számla "
            "nem fogadható be, könyvelésre és kifizetésre nem továbbítható, és haladéktalanul vissza kell utasítani a rendszer által "
            "generált indokolással. A fizetési határidő csak a hibátlan, hiánytalan és befogadott számla nyilvántartási befogadási napján kezdődhet. "
            + (f"Tervezői szerződés esetén a fizetési határidő {DESIGNER_PAYMENT_TERM_DAYS} naptári nap." if ct == DESIGN_SUBCONTRACTOR_TYPE else f"Kivitelezői alvállalkozói szerződés esetén a fizetési határidő {EXECUTION_SUBCONTRACTOR_PAYMENT_TERM_DAYS} naptári nap.")
        )
        target.insert_paragraph_before(invoice_clause)
        audit.append({"field": "mandatory_invoice_acceptance_clause", "changed": True, "target": f"before_paragraph:{insertion_index}"})

    doc.save(str(output_path))
    return audit


PLACEHOLDER_REGEXES = [
    re.compile(r"…{2,}"), re.compile(r"_{4,}"), re.compile(r"\.\s*\.\s*\.\s*\."),
    re.compile(r"2024\s*[.…_]"), re.compile(r"\bNév:\s*$", re.I), re.compile(r"\bCím:\s*$", re.I),
]


def scan_unresolved_placeholders(path: Path) -> list[str]:
    doc = create_docx_document(str(path))
    findings: list[str] = []
    texts: list[tuple[str, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        texts.append((f"paragraph:{idx}", p.text))
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                texts.append((f"table:{ti}:row:{ri}:cell:{ci}", cell.text))
    for loc, text in texts:
        compact = text.strip()
        if any(rx.search(compact) for rx in PLACEHOLDER_REGEXES):
            findings.append(f"{loc}: {compact[:180]}")
    return findings


def _configure_doc(doc: DocxDocument) -> None:
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Arial"


def _compact_landscape(doc: DocxDocument, normal_size: float = 8.5) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.42))
    doc.styles["Normal"].font.size = Pt(normal_size)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _set_cell_font_size(cell, size_pt: float) -> None:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size_pt)


def generate_data_annex(data: dict[str, Any], path: Path) -> None:
    doc = create_docx_document(); _configure_doc(doc); _compact_landscape(doc, 8.2)
    title = doc.add_paragraph("1. SZÁMÚ MELLÉKLET – KÖTELEZŐ SZERZŐDÉSES ADATLAP")
    title.style = doc.styles["Title"]; title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [
        ("Szerződésszám", data["contract_number"]), ("Szerződés kelte és helye", f"{data['contract_place']}, {data['contract_date']}"),
        *[(key, data["ids"][key]) for key in REQUIRED_IDS],
        ("Imperial szerződő cég", _party_description(data["internal_entity"])),
        ("Partner / megrendelő", _party_description(data["counterparty"])),
        ("Projekt", data["project"]["name"]), ("Helyszín", f"{data['project']['site_address']}, {data['project']['parcel_number']} hrsz."),
        ("Műszaki tartalom", data["project"]["scope"]), ("Kezdés", data["schedule"]["start_date"]),
        ("Véghatáridő", data["schedule"]["deadline"]),
        ("Nettó díj", f"{data['commercial']['net_price']} {data['commercial']['currency']}"),
        ("ÁFA", f"{data['commercial']['vat_percent']}% / {data['commercial']['vat_amount']} {data['commercial']['currency']}"),
        ("Bruttó díj", f"{data['commercial']['gross_price']} {data['commercial']['currency']}"),
        ("Számlafizetési határidő", (
            f"{data.get('invoice_controls', {}).get('payment_term_days')} naptári nap"
            if data["contract_type"] == DESIGN_SUBCONTRACTOR_TYPE else
            f"{data.get('subcontractor_controls', {}).get('payment_term_days')} naptári nap"
            if data["contract_type"] == EXECUTION_SUBCONTRACTOR_TYPE else
            "A szerződés fizetési üteme szerint"
        )),
        ("Fizetési szabály döntési alapja", OWNER_POLICY_VERSION if data["contract_type"] in SUBCONTRACTOR_TYPES else "Nem alkalmazandó"),
        ("Postai kézbesítés", f"Kötelező, igazolható mód: {data['delivery_requirements']['postal_method']}"),
        ("Elektronikus kézbesítés", "Kötelező, a partner hivatalos e-mail-címére, csatolmány-hash és üzenetazonosító naplózásával"),
        ("Mellékletek", _annex_names(data)),
    ]
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
    table.rows[0].cells[0].text = "Mező"; table.rows[0].cells[1].text = "Kitöltött érték"
    for label, value in rows:
        cells = table.add_row().cells; cells[0].text = str(label); cells[1].text = str(value)
    doc.add_heading("Fizetési ütemezés", level=1)
    pay = doc.add_table(rows=1, cols=3); pay.style = "Table Grid"
    for i, h in enumerate(("Mérföldkő", "Arány", "Esedékességi szabály")): pay.rows[0].cells[i].text = h
    for row in data["commercial"]["payment_schedule"]:
        cells = pay.add_row().cells; cells[0].text = row["milestone"]; cells[1].text = f"{row['percent']}%"; cells[2].text = row["due_rule"]
    for tbl in doc.tables:
        for row in tbl.rows:
            _prevent_row_split(row)
            for cell in row.cells: _set_cell_font_size(cell, 7.7)
    doc.save(str(path))


def generate_deduction_annex(data: dict[str, Any], path: Path) -> None:
    doc = create_docx_document(); _configure_doc(doc); _compact_landscape(doc, 7.8)
    title = doc.add_paragraph("2. SZÁMÚ MELLÉKLET – MŰSZAKI MEGFELELŐSÉG ÉS LEVONÁSI TÁBLA")
    title.style = doc.styles["Title"]; title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=5); table.style = "Table Grid"
    for i, h in enumerate(("Hiba / hiányosság", "Mérték", "Minimum összeg", "Vetítési alap", "Feloldás feltétele")): table.rows[0].cells[i].text = h
    for row in data["subcontractor_controls"]["defect_deduction_schedule"]:
        cells = table.add_row().cells
        cells[0].text = str(row.get("defect")); cells[1].text = f"{row['percent']}%" if row.get("percent") is not None else str(row.get("amount"))
        cells[2].text = str(row.get("minimum_amount", "")); cells[3].text = str(row.get("basis", "")); cells[4].text = str(row.get("release_condition", ""))
    for row in table.rows:
        _prevent_row_split(row)
        for cell in row.cells: _set_cell_font_size(cell, 7.2)
    doc.save(str(path))


def generate_checklist_doc(data: dict[str, Any], path: Path) -> None:
    doc = create_docx_document(); _configure_doc(doc)
    is_subcontractor = data["contract_type"] in {EXECUTION_SUBCONTRACTOR_TYPE, DESIGN_SUBCONTRACTOR_TYPE}
    title_text = (
        "JÓVÁHAGYÁSI, KÉZBESÍTÉSI, INDÍTÁSI ÉS SZÁMLAKAPUK"
        if is_subcontractor else
        "JÓVÁHAGYÁSI, KÉZBESÍTÉSI ÉS INDÍTÁSI KAPUK"
    )
    title = doc.add_paragraph(title_text)
    title.style = doc.styles["Title"]; title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sections = {
        "Aláírásra küldés előtt": [
            "Minden szerződésmező kitöltött; nulla helyőrző vagy üres kötelező mező",
            "Partner- és Imperial-cégadatok, személyes azonosítóadatok, projektcím, helyrajzi szám, nettó–ÁFA–bruttó ár és határidők ellenőrizve",
            "Master SHA-256 egyezés, kötelező mellékletek és jóváhagyások rögzítve",
        ],
        "Aláírás után – kötelező kettős kézbesítés": [
            "Aláírt eredeti példány igazolható postai/futárküldeményként feladva",
            "Nyomkövetési szám és feladási/kézbesítési bizonylat a ContractID-hoz mentve",
            "Azonos aláírt dokumentum elektronikusan is elküldve a partner hivatalos e-mail-címére",
            "E-mail MessageID, címzett, küldési idő és csatolmány SHA-256 rögzítve",
        ],
        "Munka megkezdése előtt": [
            "Mindkét fél által aláírt szerződés és mellékletek visszamentve",
            "Kettős kézbesítési kapu teljesült",
            "WorkStartAllowed = true",
        ],
    }
    if is_subcontractor:
        certificate_label = (
            "tervezői teljesítésigazolás"
            if data["contract_type"] == DESIGN_SUBCONTRACTOR_TYPE else
            "TIG"
        )
        sections["Számlabefogadás előtt"] = [
            f"Aláírt és elfogadott, jogosult személy által jóváhagyott {certificate_label} rendelkezésre áll",
            "Számla sorszáma, kelte, teljesítési időpontja, esedékessége, felek neve/címe/adószáma és tételsorai hiánytalanok",
            f"Nettó, ÁFA és bruttó összeg számtanilag helyes és megegyezik a {certificate_label} összegével",
            "Számla olvasható, sértetlen, nem duplikált, és a szerződés/projekt hivatkozása helyes",
            "Hiba esetén azonnali REJECTED_IMMEDIATELY státusz és jogi szövegű visszautasító értesítés",
        ]
    for heading, items in sections.items():
        doc.add_heading(heading, level=1)
        for item in items: doc.add_paragraph(item, style="List Bullet")
    doc.save(str(path))


def generate_invoice_rejection_doc(notice: dict[str, Any], path: Path) -> None:
    doc = create_docx_document(); _configure_doc(doc)
    doc.add_heading(notice["subject"], 0)
    for block in notice["body"].split("\n\n"):
        doc.add_paragraph(block)
    doc.add_heading("Jogszabályi hivatkozások", level=1)
    for text in notice["legal_basis"].values(): doc.add_paragraph(text, style="List Bullet")
    doc.save(str(path))


def generate_package(data: dict[str, Any], registry_path: Path, templates_dir: Path, output_dir: Path) -> dict[str, Any]:
    assert_valid(data)
    registry = json.loads(registry_path.read_text(encoding="utf-8"))
    entry = registry[data["contract_type"]]
    template_path = templates_dir / entry["file_name"]
    if not template_path.exists(): raise ContractValidationError(f"Template not found: {template_path}")
    actual_hash = sha256_file(template_path)
    if actual_hash != entry["sha256"]: raise ContractValidationError("Template SHA-256 mismatch; generation blocked.")
    output_dir.mkdir(parents=True, exist_ok=True)
    contract_path = output_dir / "00_contract_filled.docx"
    audit = fill_designated_fields(template_path, contract_path, data)
    unresolved = scan_unresolved_placeholders(contract_path)
    audit_failures = [item for item in audit if item.get("changed") is not True]
    generate_data_annex(data, output_dir / "01_contract_data_and_terms.docx")
    if data["contract_type"] == EXECUTION_SUBCONTRACTOR_TYPE:
        generate_deduction_annex(data, output_dir / "02_technical_compliance_deductions.docx")
        generate_checklist_doc(data, output_dir / "03_approval_delivery_and_gate_checklist.docx")
    else:
        generate_checklist_doc(data, output_dir / "02_approval_delivery_and_gate_checklist.docx")

    owner_policy = entry.get("owner_approved_policy", {})
    signing_status = "READY_FOR_APPROVAL"
    if unresolved or audit_failures:
        signing_status = "BLOCKED_UNRESOLVED_CONTRACT_FIELDS"

    fields_complete = not unresolved and not audit_failures
    status = copy.deepcopy(data.get("status", {})); status["master_hash_verified"] = True; status["all_fields_complete"] = fields_complete
    gate_data = {**data, "status": status}
    manifest = {
        "generator_version": GENERATOR_VERSION,
        "generated_at": datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "contract_type": data["contract_type"], "contract_number": data["contract_number"],
        "template": {"drive_file_id": entry["drive_file_id"], "file_name": entry["file_name"], "sha256": actual_hash},
        "ids": data["ids"], "field_fill_audit": audit, "field_fill_failures": audit_failures,
        "unresolved_placeholders": unresolved, "all_fields_complete": fields_complete,
        "dispatch_gate": dispatch_gate(data), "work_start_gate": work_start_gate(gate_data),
        "owner_approved_policy": owner_policy,
        "signing_queue_status": signing_status, "files": [],
    }
    for p in sorted(output_dir.glob("*.docx")): manifest["files"].append({"name": p.name, "sha256": sha256_file(p)})
    (output_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "source_input.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "validation_result.json").write_text(json.dumps({"valid": True, "issues": []}, ensure_ascii=False, indent=2), encoding="utf-8")
    zip_path = output_dir.with_suffix(".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(output_dir.iterdir()): zf.write(p, arcname=p.name)
    return {"output_dir": str(output_dir), "zip_path": str(zip_path), "manifest": manifest}


def load_json(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def cli(argv: Iterable[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=f"Imperial Contract Generator v{GENERATOR_VERSION}")
    sub = parser.add_subparsers(dest="command", required=True)
    p = sub.add_parser("validate"); p.add_argument("--input", required=True)
    p = sub.add_parser("generate"); p.add_argument("--input", required=True); p.add_argument("--registry", default="config/templates.json"); p.add_argument("--templates", default="master_templates"); p.add_argument("--output", required=True)
    p = sub.add_parser("invoice-gate"); p.add_argument("--input", required=True); p.add_argument("--invoice", required=True); p.add_argument("--rejection-docx")
    args = parser.parse_args(list(argv) if argv is not None else None)
    try:
        data = load_json(args.input)
        if args.command == "validate":
            issues = validate_contract(data); valid = not any(i.blocking for i in issues)
            print(json.dumps({"valid": valid, "issues": [i.as_dict() for i in issues]}, ensure_ascii=False, indent=2)); return 0 if valid else 2
        if args.command == "generate":
            print(json.dumps(generate_package(data, Path(args.registry), Path(args.templates), Path(args.output)), ensure_ascii=False, indent=2)); return 0
        invoice = load_json(args.invoice); result = invoice_acceptance_gate(data, invoice)
        if not result["accepted"] and args.rejection_docx:
            generate_invoice_rejection_doc(result["rejection_notice"], Path(args.rejection_docx))
            result["rejection_docx"] = args.rejection_docx
        print(json.dumps(result, ensure_ascii=False, indent=2)); return 0 if result["accepted"] else 3
    except (ContractValidationError, OSError, json.JSONDecodeError) as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False, indent=2), file=sys.stderr); return 1


if __name__ == "__main__":
    raise SystemExit(cli())
