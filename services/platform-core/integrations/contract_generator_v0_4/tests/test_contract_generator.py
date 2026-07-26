import copy
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from imperial_contract_generator.core import (
    ContractValidationError,
    dispatch_gate,
    generate_invoice_rejection_doc,
    generate_package,
    invoice_acceptance_gate,
    scan_unresolved_placeholders,
    select_contract_type,
    sha256_file,
    validate_contract,
    work_start_gate,
)

ROOT = Path(__file__).resolve().parents[1]

def load(name):
    return json.loads((ROOT / "examples" / name).read_text(encoding="utf-8"))

EXEC = load("subcontractor_execution_valid.json")
DESIGN = load("subcontractor_design_valid.json")
CUSTOMER_CONSTRUCTION = load("customer_construction_valid.json")
CUSTOMER_DESIGN = load("customer_design_execution_plans_valid.json")
CUSTOMER_TYPE = load("customer_type_house_design_build_valid.json")
ALL_CONTRACTS = [EXEC, DESIGN, CUSTOMER_CONSTRUCTION, CUSTOMER_DESIGN, CUSTOMER_TYPE]


def complete_dispatch(data):
    d = copy.deepcopy(data)
    h = "b" * 64
    d["dispatch_status"] = {
        "internal_signed_original_present": True,
        "internal_signature_date": "2026-07-18",
        "signed_document_sha256": h,
        "postal": {
            "sent": True,
            "sent_at": "2026-07-18T13:00:00+02:00",
            "recipient_address": d["counterparty"]["postal_address"],
            "original_copy_count": 1,
            "tracking_number": "RR123456789HU",
            "proof_file_id": "DRV-POST-001",
        },
        "electronic": {
            "sent": True,
            "sent_at": "2026-07-18T13:05:00+02:00",
            "recipient_email": d["counterparty"]["email"],
            "message_id": "<msg-001@example.com>",
            "attachment_sha256": h,
        },
    }
    return d


class ContractGeneratorTests(unittest.TestCase):
    def test_contract_type_selection_all_five(self):
        cases = [
            ({"relationship": "partner", "service": "construction"}, "subcontractor_execution"),
            ({"relationship": "partner", "service": "design"}, "subcontractor_design"),
            ({"relationship": "customer", "service": "construction"}, "customer_construction"),
            ({"relationship": "customer", "service": "design", "execution_plans": True}, "customer_design_execution_plans"),
            ({"relationship": "customer", "service": "design_build", "type_house": True}, "customer_type_house_design_build"),
        ]
        for data, expected in cases:
            with self.subTest(expected=expected):
                self.assertEqual(expected, select_contract_type(data))

    def test_all_five_examples_are_valid(self):
        for contract in ALL_CONTRACTS:
            with self.subTest(contract_type=contract["contract_type"]):
                self.assertEqual([], [i.as_dict() for i in validate_contract(copy.deepcopy(contract)) if i.blocking])

    def test_contract_header_fields_are_mandatory(self):
        for key in ("contract_number", "contract_date", "contract_place"):
            data = copy.deepcopy(EXEC); data.pop(key)
            codes = {i.code for i in validate_contract(data)}
            self.assertIn("CONTRACT_HEADER_MISSING", codes)

    def test_all_crm_ids_are_mandatory(self):
        for key in ("CompanyID", "PersonID", "OpportunityID", "ProjectID", "PartnerID"):
            data = copy.deepcopy(EXEC); data["ids"][key] = ""
            self.assertIn("ID_MISSING", {i.code for i in validate_contract(data)})

    def test_every_company_field_is_mandatory(self):
        fields = ("name", "address", "postal_address", "email", "phone", "short_name", "registration_number",
                  "tax_number", "registered_office", "bank_account", "representative", "representative_title")
        for prefix in ("internal_entity", "counterparty"):
            for field in fields:
                data = copy.deepcopy(EXEC); data[prefix][field] = ""
                codes = {i.code for i in validate_contract(data)}
                self.assertTrue({"PARTY_DATA_MISSING", "COMPANY_DATA_MISSING"} & codes, (prefix, field, codes))

    def test_every_natural_person_field_is_mandatory(self):
        fields = ("name", "address", "postal_address", "email", "phone", "birth_place", "birth_date", "mother_name",
                  "identity_document_type", "identity_document_number")
        for field in fields:
            data = copy.deepcopy(CUSTOMER_CONSTRUCTION); data["counterparty"][field] = ""
            codes = {i.code for i in validate_contract(data)}
            self.assertTrue({"PARTY_DATA_MISSING", "PERSONAL_DATA_MISSING"} & codes, (field, codes))

    def test_project_address_parcel_scope_and_price_are_mandatory(self):
        for path in ("site_address", "parcel_number", "scope"):
            data = copy.deepcopy(EXEC); data["project"][path] = ""
            self.assertIn("PROJECT_DATA_MISSING", {i.code for i in validate_contract(data)})
        for field in ("net_price", "vat_percent", "vat_amount", "gross_price", "currency", "payment_schedule"):
            data = copy.deepcopy(EXEC); data["commercial"][field] = "" if field != "payment_schedule" else []
            self.assertIn("COMMERCIAL_DATA_MISSING", {i.code for i in validate_contract(data)})

    def test_price_arithmetic_and_payment_schedule_must_be_exact(self):
        data = copy.deepcopy(EXEC); data["commercial"]["gross_price"] += 100
        self.assertIn("GROSS_CALCULATION_MISMATCH", {i.code for i in validate_contract(data)})
        data = copy.deepcopy(EXEC); data["commercial"]["payment_schedule"][0]["percent"] = 19
        self.assertIn("PAYMENT_SCHEDULE_TOTAL_INVALID", {i.code for i in validate_contract(data)})

    def test_dual_delivery_configuration_is_mandatory(self):
        for field, code in (("postal_required", "POSTAL_DELIVERY_RULE_MISSING"), ("electronic_required", "ELECTRONIC_DELIVERY_RULE_MISSING")):
            data = copy.deepcopy(EXEC); data["delivery_requirements"][field] = False
            self.assertIn(code, {i.code for i in validate_contract(data)})
        data = copy.deepcopy(EXEC); data["delivery_requirements"]["postal_original_copies"] = 0
        self.assertIn("POSTAL_ORIGINAL_COPY_MISSING", {i.code for i in validate_contract(data)})

    def test_execution_only_controls(self):
        mutations = [
            (lambda d: d["subcontractor_controls"].update(warranty_retention_percent=4.99), "RETENTION_TOO_LOW"),
            (lambda d: d["subcontractor_controls"].update(client_services_deduction_percent=0.5), "SERVICE_FEE_TOO_LOW"),
            (lambda d: d["subcontractor_controls"].update(payment_term_days=29), "EXECUTION_PAYMENT_TERM_INVALID"),
            (lambda d: d["subcontractor_controls"].update(payment_term_days=31), "EXECUTION_PAYMENT_TERM_INVALID"),
            (lambda d: d["subcontractor_controls"].update(defect_deduction_schedule=[]), "DEDUCTION_SCHEDULE_MISSING"),
            (lambda d: d["subcontractor_controls"]["invoice_gate"].update(upstream_customer_tig_required=False), "UPSTREAM_TIG_RULE_MISSING"),
        ]
        for mutate, code in mutations:
            data = copy.deepcopy(EXEC); mutate(data)
            self.assertIn(code, {i.code for i in validate_contract(data)})

    def test_owner_approved_payment_terms_are_exact(self):
        data = copy.deepcopy(DESIGN); data["invoice_controls"]["payment_term_days"] = 7
        self.assertIn("DESIGN_PAYMENT_TERM_INVALID", {i.code for i in validate_contract(data)})
        data = copy.deepcopy(DESIGN); data["invoice_controls"]["payment_term_days"] = 9
        self.assertIn("DESIGN_PAYMENT_TERM_INVALID", {i.code for i in validate_contract(data)})
        self.assertEqual(8, DESIGN["invoice_controls"]["payment_term_days"])
        self.assertEqual(30, EXEC["subcontractor_controls"]["payment_term_days"])

    def test_designer_profile_excludes_construction_controls(self):
        codes = {i.code for i in validate_contract(copy.deepcopy(DESIGN))}
        self.assertNotIn("RETENTION_TOO_LOW", codes)
        self.assertNotIn("UPSTREAM_TIG_RULE_MISSING", codes)
        data = copy.deepcopy(DESIGN); data["subcontractor_controls"] = copy.deepcopy(EXEC["subcontractor_controls"])
        issues = validate_contract(data)
        self.assertEqual([], [i for i in issues if i.blocking])
        self.assertIn("CONSTRUCTION_CONTROLS_NOT_APPLICABLE", {i.code for i in issues})

    def test_dispatch_gate_requires_trackable_post_and_electronic_proof(self):
        result = dispatch_gate(copy.deepcopy(EXEC))
        self.assertFalse(result["allowed"])
        self.assertIn("SIGNED_ORIGINAL_NOT_POSTED", result["blockers"])
        self.assertIn("SIGNED_CONTRACT_NOT_SENT_ELECTRONICALLY", result["blockers"])
        self.assertTrue(dispatch_gate(complete_dispatch(EXEC))["allowed"])

    def test_dispatch_gate_checks_recipient_and_hash(self):
        data = complete_dispatch(EXEC)
        data["dispatch_status"]["electronic"]["recipient_email"] = "wrong@example.com"
        data["dispatch_status"]["electronic"]["attachment_sha256"] = "c" * 64
        blockers = dispatch_gate(data)["blockers"]
        self.assertIn("ELECTRONIC_RECIPIENT_MISMATCH", blockers)
        self.assertIn("ELECTRONIC_ATTACHMENT_HASH_MISMATCH", blockers)

    def test_work_start_requires_both_signatures_and_both_delivery_channels(self):
        data = complete_dispatch(EXEC)
        data["status"].update(contract_status="SIGNED", signed_contract_present=True, both_parties_signed=True,
                              signed_contract_file_id="DRV-SIGNED-001", master_hash_verified=True,
                              all_required_annexes_present=True, all_fields_complete=True)
        self.assertTrue(work_start_gate(data)["allowed"])
        data["status"]["both_parties_signed"] = False
        self.assertIn("BOTH_PARTIES_SIGNATURE_MISSING", work_start_gate(data)["blockers"])

    def test_execution_invoice_accepted_only_with_perfect_invoice_and_accepted_tig(self):
        invoice = load("invoice_execution_valid.json")
        result = invoice_acceptance_gate(copy.deepcopy(EXEC), invoice)
        self.assertTrue(result["accepted"])
        self.assertEqual("50000.00", result["calculation"]["warranty_retention"])
        self.assertEqual("10000.00", result["calculation"]["client_services_deduction"])
        self.assertEqual("50000.00", result["calculation"]["good_performance_guarantee"])
        self.assertEqual("890000.00", result["calculation"]["standard_payable"])
        self.assertEqual("2026-10-20", result["calculation"]["standard_due_date"])

    def test_designer_invoice_requires_designer_performance_certificate_without_construction_deductions(self):
        result = invoice_acceptance_gate(copy.deepcopy(DESIGN), load("invoice_design_valid.json"))
        self.assertTrue(result["accepted"])
        self.assertEqual("0.00", result["calculation"]["warranty_retention"])
        self.assertEqual("0.00", result["calculation"]["client_services_deduction"])
        self.assertEqual("533400.00", result["calculation"]["standard_payable"])
        self.assertEqual("2026-11-08", result["calculation"]["standard_due_date"])

    def test_invoice_without_accepted_signed_tig_is_rejected_immediately(self):
        invoice = load("invoice_execution_valid.json")
        invoice["performance_certificate"]["status"] = "DRAFT"
        invoice["performance_certificate"]["authorized_acceptor_signed"] = False
        result = invoice_acceptance_gate(copy.deepcopy(EXEC), invoice)
        self.assertEqual("REJECTED_IMMEDIATELY", result["status"])
        self.assertIn("TIG_NOT_ACCEPTED", result["blockers"])
        self.assertIn("TIG_AUTHORIZED_SIGNATURE_MISSING", result["blockers"])

    def test_formal_or_content_invoice_error_is_rejected_immediately_with_legal_notice(self):
        result = invoice_acceptance_gate(copy.deepcopy(EXEC), load("invoice_execution_invalid.json"))
        self.assertFalse(result["accepted"])
        self.assertEqual("REJECTED_IMMEDIATELY", result["status"])
        notice = result["rejection_notice"]
        self.assertIn("nem fogadjuk be", notice["body"])
        self.assertIn("Áfa tv. 168/A. §", notice["body"])
        self.assertIn("Áfa tv. 169. §", notice["body"])
        self.assertIn("Áfa tv. 170. §", notice["body"])
        self.assertNotIn("supplier tax mismatch", notice["body"])

    def test_line_item_sum_error_is_rejected(self):
        invoice = load("invoice_execution_valid.json")
        invoice["line_items"][0]["net_amount"] -= 1000
        result = invoice_acceptance_gate(copy.deepcopy(EXEC), invoice)
        self.assertIn("INVOICE_LINE_NET_MISMATCH", result["blockers"])
        self.assertIn("INVOICE_LINE_SUM_NET_MISMATCH", result["blockers"])

    def test_invoice_gate_is_not_available_for_customer_contracts(self):
        with self.assertRaises(ContractValidationError):
            invoice_acceptance_gate(copy.deepcopy(CUSTOMER_CONSTRUCTION), load("invoice_execution_valid.json"))

    def test_rejection_notice_docx_is_generated(self):
        result = invoice_acceptance_gate(copy.deepcopy(EXEC), load("invoice_execution_invalid.json"))
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "rejection.docx"
            generate_invoice_rejection_doc(result["rejection_notice"], path)
            self.assertTrue(path.exists())
            text = "\n".join(p.text for p in Document(path).paragraphs)
            self.assertIn("Számla befogadásának elutasítása", text)
            self.assertIn("nem továbbítjuk", text)

    def test_template_hash_registry(self):
        registry = json.loads((ROOT / "config/templates.json").read_text(encoding="utf-8"))
        for entry in registry.values():
            self.assertEqual(entry["sha256"], sha256_file(ROOT / "master_templates" / entry["file_name"]))

    def test_all_five_contracts_generate_without_placeholders_or_failed_field_writes(self):
        registry = ROOT / "config/templates.json"
        templates = ROOT / "master_templates"
        with tempfile.TemporaryDirectory() as tmp:
            for contract in ALL_CONTRACTS:
                with self.subTest(contract_type=contract["contract_type"]):
                    out = Path(tmp) / contract["contract_type"]
                    result = generate_package(copy.deepcopy(contract), registry, templates, out)
                    self.assertTrue(Path(result["zip_path"]).exists())
                    manifest = result["manifest"]
                    self.assertEqual([], manifest["unresolved_placeholders"])
                    self.assertEqual([], manifest["field_fill_failures"])
                    self.assertTrue(manifest["all_fields_complete"])
                    self.assertIn("owner_approved_policy", manifest)
                    self.assertNotEqual("BLOCKED_LEGAL_TEMPLATE_CONFLICT", manifest["signing_queue_status"])
                    self.assertEqual([], scan_unresolved_placeholders(out / "00_contract_filled.docx"))
                    self.assertTrue((out / "01_contract_data_and_terms.docx").exists())
                    self.assertTrue((out / "02_approval_delivery_and_gate_checklist.docx").exists() or
                                    (out / "03_approval_delivery_and_gate_checklist.docx").exists())

    def test_execution_and_design_annex_sets_are_separate(self):
        with tempfile.TemporaryDirectory() as tmp:
            eout = Path(tmp) / "exec"; dout = Path(tmp) / "design"
            generate_package(copy.deepcopy(EXEC), ROOT / "config/templates.json", ROOT / "master_templates", eout)
            generate_package(copy.deepcopy(DESIGN), ROOT / "config/templates.json", ROOT / "master_templates", dout)
            self.assertTrue((eout / "02_technical_compliance_deductions.docx").exists())
            self.assertFalse((dout / "02_technical_compliance_deductions.docx").exists())
            self.assertEqual("0.4.0", json.loads((dout / "manifest.json").read_text(encoding="utf-8"))["generator_version"])

    def test_generated_designer_contract_has_no_review_highlights_or_red_instructions(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "design"
            generate_package(copy.deepcopy(DESIGN), ROOT / "config/templates.json", ROOT / "master_templates", out)
            doc = Document(out / "00_contract_filled.docx")
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertNotIn("a megfelelő maradjon benne", text)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    self.assertIsNone(run.font.highlight_color)
                    if run.font.color and run.font.color.rgb:
                        self.assertNotEqual("FF0000", str(run.font.color.rgb))

    def test_invoice_checklist_is_only_in_subcontractor_packages(self):
        with tempfile.TemporaryDirectory() as tmp:
            cout = Path(tmp) / "customer"
            dout = Path(tmp) / "designer"
            eout = Path(tmp) / "execution"
            generate_package(copy.deepcopy(CUSTOMER_CONSTRUCTION), ROOT / "config/templates.json", ROOT / "master_templates", cout)
            generate_package(copy.deepcopy(DESIGN), ROOT / "config/templates.json", ROOT / "master_templates", dout)
            generate_package(copy.deepcopy(EXEC), ROOT / "config/templates.json", ROOT / "master_templates", eout)
            customer_text = "\n".join(p.text for p in Document(cout / "02_approval_delivery_and_gate_checklist.docx").paragraphs)
            designer_text = "\n".join(p.text for p in Document(dout / "02_approval_delivery_and_gate_checklist.docx").paragraphs)
            execution_text = "\n".join(p.text for p in Document(eout / "03_approval_delivery_and_gate_checklist.docx").paragraphs)
            self.assertNotIn("Számlabefogadás előtt", customer_text)
            self.assertIn("tervezői teljesítésigazolás", designer_text)
            self.assertIn("TIG", execution_text)

    def test_execution_signature_page_has_single_non_applicable_witness_line(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "execution"
            generate_package(copy.deepcopy(EXEC), ROOT / "config/templates.json", ROOT / "master_templates", out)
            text = "\n".join(p.text for p in Document(out / "00_contract_filled.docx").paragraphs)
            marker = "Tanúk: nem alkalmazandók – mindkét fél cégszerű aláírással jár el."
            self.assertEqual(1, text.count(marker))
            self.assertNotIn("Előttük, mint tanúk előtt", text)

    def test_generated_contracts_contain_no_stale_reference_party_or_bank_data(self):
        stale = ["Budai Zsolt", "Szántó Éva", "10401804-50527078-84761000", "25322840-2-13", "2682 Püspökhatvan"]
        with tempfile.TemporaryDirectory() as tmp:
            for contract in ALL_CONTRACTS:
                out = Path(tmp) / contract["contract_type"]
                generate_package(copy.deepcopy(contract), ROOT / "config/templates.json", ROOT / "master_templates", out)
                doc = Document(out / "00_contract_filled.docx")
                text = "\n".join(p.text for p in doc.paragraphs)
                for marker in stale:
                    self.assertNotIn(marker, text, (contract["contract_type"], marker))


if __name__ == "__main__":
    unittest.main()
